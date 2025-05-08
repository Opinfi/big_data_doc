#! python3.11

#   BIG DATA DOCUMENTS =============================== PROYECTO DE MAESTRIA
#   AUTOR: OSWALDO WISTON PIN FIGUEROA ====================================
#
# Bibliotecas auxiliares ----------------------------------------------------------------------------------------
import numpy as np
import operator # Módulo estándar de Python para operadores
import matplotlib
import matplotlib.pyplot as plt
matplotlib.use('Agg')  # Configurar el backend a 'Agg' para evitar problemas con hilos
from PIL import Image
import io
import os
from datetime import datetime, timedelta
import pytz
import random
import requests
import secrets
import zipfile
import rarfile
# Bibliotecas para los procesamientos y aplicación de los modelos
# predictivos y de clasificación a las tablas extraídas de los 
# documentos, para ser enviados al lado del cliente
from prophet import Prophet
from sklearn.linear_model import LinearRegression
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.svm import SVC
from sklearn.preprocessing import LabelEncoder, OneHotEncoder
from sklearn.preprocessing import StandardScaler
from sklearn.compose import ColumnTransformer
import lightgbm as lgb
from prophet import Prophet
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import confusion_matrix, accuracy_score, classification_report
from sklearn.metrics import precision_score, recall_score, f1_score
# Modulos para poder explorar los datos estructurados de los diferentes tipos de documentos ---------------------
# Documentos Office Word
from docx import Document
from lxml import etree
from docx.shared import Inches, Pt, Cm
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.table import _Cell
from PIL import Image
import math
from docx2pdf import convert
import pythoncom  # Importar pythoncom para inicializar COM
from docx2python import docx2python
from docx2python.iterators import iter_tables
import logging
# Documentos PDF
import tabula
# Documentos Excel y mas para trabajo con Big Data
import dask.dataframe as dd
import pandas as pd
import pdfplumber
import json
# Modulo para el servidor Python - Flask API --------------------------------------------------------------------
from flask import render_template, Flask, request, url_for, session, Response
from flask import jsonify, send_from_directory, send_file
#from flask_session import Session
from markupsafe import Markup
#para manejar las base de datos y la conformacion de los datos
from sqlalchemy import text
from sqlalchemy import create_engine, MetaData, engine
from sqlalchemy.exc import SQLAlchemyError

# ----------------------------- V A R I A B L E S   Y   F U N C I O N E S ------------------------------------------

#Instanciamos con la clase Flask, usando directorios templates y static/...
app = Flask(__name__, template_folder='templates', static_folder='static')
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024 * 1024  # 20 GB
app.config['SESSION_COOKIE_SECURE'] = False  # Solo enviar cookies sobre HTTPS, de ser True
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Prevenir acceso a la cookie desde JavaScript
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # Prevenir ataques CSRF
#clave secreta  
app.secret_key = 'tu_clave_secreta_fija_y_seguraOswaldo2025***/' #secrets.token_hex(16) # Genera una clave secreta de 32 caracteres
#duracion de la sesion abierta, para el caso del administrador del sitema experimental
app.permanent_session_lifetime = timedelta(days=2)

# --------------- F U N C I O N ES   A U X I L I A R E S  ----------------
#Variables globales para la oonexion a la base de datos de preguntas/respuestas
mydb = None
conn = None
metadata = MetaData()

def connect_dbase():
    # Establecemos la conexion con la base de datos 'big_data_doc' que contiene
    # La entidad-relacion de los datos involucrados como los usuarios...
    ##host = "127.0.0.1",
    ##user = "root",
    ##password = "",
    ##database = "big_data_doc"
    ##port = "3306"
    global mydb, metadata, conn
    mydb = create_engine("mysql+mysqlconnector://root:@127.0.0.1:3306/big_data_doc?charset=utf8", 
        pool_pre_ping=True, pool_recycle=3600*24)
    metadata.reflect(bind=mydb)
    conn = mydb.connect()

#Llamamos a la funcion para conectarnos a la base de datos
connect_dbase()

def reset_global_variables():
    session['username'] = ''
    session['password'] = ''
    session['tipo_usuario'] = 'anonimo'
    session['url_foto'] = 'static/img/anonimo_blue.png'
    session.pop('first_time', None)

def set_global_variables(username: str, password: str, url_foto: str, tipo_usuario: str):
    session['username'] = username
    session['password'] = password
    session['url_foto'] = url_foto
    session['tipo_usuario'] = tipo_usuario

def get_global_variables():
    username = session.get('username')
    if username==None:
        username = ''
    password = session.get('password')
    if password==None:
        password = ''
    url_foto = session.get('url_foto')
    if url_foto==None:
        url_foto = 'static/img/anonimo_blue.png'
    tipo = session.get('tipo_usuario')
    if tipo==None:
        tipo = 'anonimo'
    gvar = {
        'username'      : username,
        'password'      : password,
        'url_foto'      : url_foto,
        'tipo_usuario'  : tipo,
    }
    return gvar

# Limpiar la ruta de carpeta static / tmp ...
# Ruta a limpiar
TMP_DIR = os.path.join('static', 'tmp')

def limpiar_directorio(ruta):
    """Elimina todos los archivos dentro del directorio dado."""
    try:
        # Verificar si la ruta existe
        if not os.path.exists(ruta):
            os.makedirs(ruta)  # Crear el directorio si no existe
            return
        # Listar y eliminar archivos uno por uno
        for archivo in os.listdir(ruta):
            archivo_path = os.path.join(ruta, archivo)
            try:
                if os.path.isfile(archivo_path):
                    os.unlink(archivo_path)  # Eliminar archivo
            except Exception as e:
                print(f"Error al eliminar {archivo_path}: {e}")
    except Exception as e:
        print(f"Error al limpiar el directorio: {e}")
# Limpiar el directorio al iniciar el servidor
limpiar_directorio(TMP_DIR)

# Funciones auxiliares para operaciones sobre las tablas...
# ----------------------------------------------------------
# Etiquetas de los procesos asociados a su texto legible
Procesos_Label2Descrip = {
	"Quitar_duplicados" : "Quitar duplicados",
	"Recuperar_num_corruptos" : "Recuperar n&uacute;meros corruptos",
	"Estadistica_descriptiva" : "Estad&iacute;stica descriptiva",
	"Frecuencias_densidad_prob" : "Frecuencias (<em>Densidad Probabil&iacute;stica</em>)",
	"Prediccion_lineal" : "Predicci&oacute;n por regresi&oacute;n lineal",
	"Prediccion_nolineal" : "Predicci&oacute;n por regresi&oacute;n no lineal",
	"Prediccion_SeriesTemporales" : "Predicci&oacute;n por series temporales",
	"Clasif_Arbol_decision" : "Clasif. &Aacute;rbol de decisi&oacute;n",
	"Clasif_Random_forest" : "Clasif. Bosque aleatorio (<em>Random forest</em>)",
	"Clasif_Regresion_logistica" : "Clasif. Regresi&oacute;n log&iacute;stica",
	"Clasif_SVM" : "Clasif. M&aacute;quina de soporte vectorial (<em>SVM</em>)",
    "Generar_grafico_dispersion" : "Gr&aacute;fico de dispersi&oacute;n",
}
def Quitar_duplicados(tabla_dict, columnas_ref):
    # Convertir el diccionario en un DataFrame de Dask
    df = dd.from_pandas(pd.DataFrame(tabla_dict), npartitions=10)  # Dividir en 10 particiones
    # Eliminar duplicados basados en las columnas listadas en columnas_ref
    df_sin_duplicados = df.drop_duplicates(subset=columnas_ref)
    # Ejecutar las operaciones y obtener el resultado
    resultado = df_sin_duplicados.compute()
    return resultado

def extraer_numeros(valor):
    if pd.isna(valor):
        return valor
    valor_str = str(valor).strip()
    # Primero verificamos si tiene comas (para separadores de miles)
    if "," in valor_str:
        # Patrón para números con comas como separadores de miles
        patron = r"[-+]?\d{1,3}(?:,\d{3})*(?:\.\d*)?"
    else:
        # Patrón para números sin comas (no forzar grupos de 3)
        patron = r"[-+]?\d+(?:\.\d*)?|[-+]?\.\d+"
    match = re.search(patron, valor_str)    
    if match:
        num_str = match.group(0).replace(",", "")  # Eliminar comas
        try:
            return float(num_str) if "." in num_str else int(num_str)
        except (ValueError, TypeError):
            return None
    return None

def RecuperarNumCorruptos(tabla, columnas):
    """
    Limpia y extrae números de una o mas columnas de un DataFrame.
    :param df: DataFrame de Pandas.
    :param columna: Nombre de la columna a limpiar.
    :return: Serie con los números limpios.
    """
    # Convertir el diccionario en un DataFrame de Dask
    df = pd.DataFrame(tabla)  # Dividir en 3 particiones
    # Aplicar la función a cada columna...
    for columna in columnas:
        df[columna] = df[columna].apply(extraer_numeros)  # ¡Asignar el resultado!
    return df

def EstadisticaDescriptiva(tabla, columnas):
    """
    Calcula estadísticas descriptivas para columnas numéricas de un diccionario.
    Argumentos:
        tabla (dict): Diccionario con los datos.
        columnas (list): Lista de columnas a analizar.
    Retorna:
        DataFrame: Estadísticas descriptivas por columna.
    """
    # Convertir el diccionario a DataFrame y seleccionar columnas
    df = pd.DataFrame(tabla)
    df_seleccionado = df.loc[:, columnas].copy()  # Usar .loc y .copy() para evitar la advertencia
    # Convertir a numérico (manejar errores con 'coerce')
    df_seleccionado = df_seleccionado.apply(pd.to_numeric, errors='coerce')
    # Calcular estadísticas
    estadisticas = {
        'Máximo': df_seleccionado.max(),
        'Mínimo': df_seleccionado.min(),
        'Media': df_seleccionado.mean(),
        'Desviación Estándar': df_seleccionado.std(),
        'Varianza': df_seleccionado.var(),
        'Mediana': df_seleccionado.median(),
        'Primer Cuartil (Q1)': df_seleccionado.quantile(0.25),
        'Tercer Cuartil (Q3)': df_seleccionado.quantile(0.75),
        'Rango Intercuartílico (IQR)': df_seleccionado.quantile(0.75) - df_seleccionado.quantile(0.25)
    }
    # Crear DataFrame y transponer
    df_estadisticas = pd.DataFrame(estadisticas)
    df_estadisticas = df_estadisticas.T  # Transponer: estadísticos como filas
    df_estadisticas.reset_index(inplace=True)
    df_estadisticas.rename(columns={'index': 'Estadístico'}, inplace=True)
    return df_estadisticas

def Prediccion_lineal(x, y, colX, colY, epocas):
	# Datos históricos (x e y)
	x_historico = np.array(x)  # x no está ordenado y tiene valores repetidos
	y_historico = np.array(y)  # y correspondiente
	# Ordenar x e y según x
	sorted_indices = np.argsort(x_historico)
	x_sorted = x_historico[sorted_indices]
	y_sorted = y_historico[sorted_indices]
	# Entrenar el modelo de regresión lineal
	modelo = LinearRegression()
	# reshape para convertir x en una matriz 2D
	modelo.fit(x_sorted.reshape(-1, 1), y_sorted)
	# Predecir valores futuros
	x_futuro = np.linspace(min(x_sorted), max(x_sorted) + epocas, 200).reshape(-1, 1)  # Extender epocas unidades más
	y_pred = modelo.predict(x_futuro)
	# Convertir a un DataFrame los resultados que estan en numpy.array
	df = pd.DataFrame({
		colX: x_futuro.flatten(),
		colY: y_pred.flatten()
	})
	return df

def Prediccion_nolineal(x, y, colX, colY, epocas):
    # Preprocesamiento: eliminar duplicados y ordenar
    df_aux = pd.DataFrame({'x': np.array(x), 'y': np.array(y)})
    df_aux = df_aux.groupby('x').mean().reset_index()
    x_sorted = df_aux['x'].values
    y_sorted = df_aux['y'].values
    # Normalización (opcional pero recomendado para LightGBM)
    scaler_x = StandardScaler()
    x_scaled = scaler_x.fit_transform(x_sorted.reshape(-1, 1))
    # Entrenamiento con validación
    X_train, X_val, y_train, y_val = train_test_split(
        x_scaled, y_sorted, test_size=0.2, random_state=42
    )
    # Configuración robusta del modelo
    modelo = lgb.LGBMRegressor(
        n_estimators=2000,
        learning_rate=0.001,
        max_depth=10,
        num_leaves=63,
        min_child_samples=5,
        reg_alpha=0.1,
        reg_lambda=0.1,
        early_stopping_round=100,
        verbosity=-1
    )
    # Entrenar con nombres de características
    X_train_df = pd.DataFrame(X_train, columns=['feature'])
    modelo.fit(
        X_train_df, y_train,
        eval_set=[(pd.DataFrame(X_val, columns=['feature']), y_val)],
        eval_metric='rmse'
    )
    # Predicción
    x_futuro = np.linspace(min(x_sorted), max(x_sorted) + epocas, 100)
    x_futuro_scaled = scaler_x.transform(x_futuro.reshape(-1, 1))
    y_pred = modelo.predict(pd.DataFrame(x_futuro_scaled, columns=['feature']))
    # Retornando solucion...
    return pd.DataFrame({colX: x_futuro, colY: y_pred})

def detectar_frecuencia(df, colX_ds):
    """
    Detecta la frecuencia temporal (anual, mensual, etc.) y si el formato es numérico.
    Retorna:
        - freq (str): Frecuencia para Prophet ('D', 'M', 'YE', etc.).
        - es_numerico (bool): Si el formato original es numérico (ej: años como 1997).
    """
    # Verificar si la columna es numérica (ej: años como 1997)
    es_numerico = pd.api.types.is_numeric_dtype(df[colX_ds])
    # Convertir a datetime
    try:
        if es_numerico:
            df[colX_ds] = pd.to_datetime(df[colX_ds], format='%Y')  # Años numéricos
        else:
            df[colX_ds] = pd.to_datetime(df[colX_ds])  # Fechas estándar
    except:
        raise ValueError("Formato temporal no reconocido. Use años (1997) o fechas (YYYY-MM-DD).")
    # Calcular diferencias entre fechas
    df = df.sort_values(colX_ds).reset_index(drop=True)
    diffs = df[colX_ds].diff().dropna()
    if len(diffs) == 0:
        return 'D', es_numerico  # Valor por defecto (diario)
    diff_mas_comun = diffs.mode()[0]
    # Mapear diferencias a frecuencias de Prophet
    if diff_mas_comun <= timedelta(minutes=1):
        return 'T', es_numerico  # Minutos
    elif diff_mas_comun <= timedelta(hours=1):
        return 'H', es_numerico  # Horas
    elif diff_mas_comun <= timedelta(days=1):
        return 'D', es_numerico  # Diario
    elif diff_mas_comun <= timedelta(days=7):
        return 'W', es_numerico  # Semanal
    elif diff_mas_comun <= timedelta(days=31):
        return 'M', es_numerico  # Mensual
    elif diff_mas_comun < timedelta(days=365):
        return 'Q', es_numerico  # Trimestral
    else:
        return 'YE', es_numerico  # Anual

def Prediccion_SeriesTemporales(df, colX_ds, colY_y, epocas, include_history=True):
    """
    Predice series temporales manteniendo el formato original y detectando automáticamente la frecuencia.
    """
    # Detectar frecuencia y formato original
    freq, es_numerico = detectar_frecuencia(df.copy(), colX_ds)
    # Preparar datos para Prophet (copiar para no modificar el original)
    df_prophet = df.copy()
    df_prophet['ds'] = pd.to_datetime(df_prophet[colX_ds], format='%Y') if es_numerico else pd.to_datetime(df_prophet[colX_ds])
    df_prophet['y'] = df_prophet[colY_y]
    # Modelado con ajustes automáticos
    modelo = Prophet(
        changepoint_prior_scale=0.05,  # Más sensible a cambios
        seasonality_mode='multiplicative',
        yearly_seasonality=True,
        weekly_seasonality=False,
        daily_seasonality=False,
        holidays=None,
        mcmc_samples=300  # Muestreo bayesiano para incertidumbre
    )
    # Añadir componentes personalizados según frecuencia
    if freq in ['A', 'YE', 'Y']:
        modelo.add_seasonality(name='yearly', period=365, fourier_order=10)
    elif freq == 'Q':
        modelo.add_seasonality(name='quarterly', period=91.25, fourier_order=5)
    # Entrenar modelo
    modelo.fit(df_prophet[['ds', 'y']])
    # Generar fechas futuras
    futuro = modelo.make_future_dataframe(periods=epocas, freq=freq, include_history=include_history)
    # Predecir
    forecast = modelo.predict(futuro)
    resultado = forecast[['ds', 'yhat']].rename(columns={'yhat': 'y'})
    # Revertir al formato original
    if es_numerico:
        resultado['ds'] = resultado['ds'].dt.year  # Extraer solo el año (ej: 2023)
    else:
        # Mantener formato datetime original (ej: '2023-01-01')
        resultado['ds'] = resultado['ds'].dt.strftime('%Y-%m-%d')  # Ajusta el formato según necesites
    # retornando solucion...
    return resultado.rename(columns={'ds': colX_ds, 'y': colY_y})

# ----------------- Modelos de clasificación ------------------------------

def evaluar_modelo(y_test, y_pred):
    """
    Evalúa un modelo de clasificación a partir de y_test y y_pred.
    Devuelve un diccionario con cuatro tablas:
    1. matriz_confusion: Matriz de confusión.
    2. metricas: Métricas principales (exactitud, precisión, recall, F1-score).
    3. reporte_clasif: Reporte detallado por clase.
    4. metricas_globales: Métricas globales (accuracy, promedios macro y ponderados).
    """
    # Generar la matriz de confusión
    conf_matrix = confusion_matrix(y_test, y_pred)
    # Obtener los nombres de las clases únicas en y_test
    clases = np.unique(y_test)
    # Construir la matriz de confusión como un diccionario
    matriz_confusion = {"Referencia vs Predicción": clases.tolist()}
    for i, clase_real in enumerate(clases):
        matriz_confusion[clase_real] = conf_matrix[i].tolist()
    # Calcular métricas principales
    accuracy = 100 * accuracy_score(y_test, y_pred)  # Exactitud
    precision = 100 * precision_score(y_test, y_pred, average='weighted')  # Precisión (ponderada)
    recall = 100 * recall_score(y_test, y_pred, average='weighted')  # Recall (ponderado)
    f1 = 100 * f1_score(y_test, y_pred, average='weighted')  # F1-Score (ponderado)
    # Construir el diccionario de métricas principales
    metricas = {
        "Métricas": ["Exactitud (Accuracy)", "Precisión (Precision)", "Recall (Sensibilidad)", "F1-Score"],
        "Valores": [accuracy, precision, recall, f1]
    }
    # Reporte de clasificación detallado
    class_report = classification_report(y_test, y_pred, output_dict=True)
    # Convertir el reporte de clasificación a un diccionario de columnas y listas
    reporte_clasif = {
        "Clase": [],
        "Precisión": [],
        "Recall": [],
        "F1-Score": [],
        "Soporte": []
    }
    for clase, valores in class_report.items():
        if clase not in ["accuracy", "macro avg", "weighted avg"]:  # Excluir métricas globales
            reporte_clasif["Clase"].append(clase)
            reporte_clasif["Precisión"].append(100 * valores["precision"])
            reporte_clasif["Recall"].append(100 * valores["recall"])
            reporte_clasif["F1-Score"].append(100 * valores["f1-score"])
            reporte_clasif["Soporte"].append(100 * valores["support"])
    # Crear una tabla diccionario para las métricas globales
    metricas_globales = {
        "Métrica": ["Exactitud (Accuracy)", "Precisión (Macro Avg)", "Recall (Macro Avg)", "F1-Score (Macro Avg)", 
                    "Precisión (Weighted Avg)", "Recall (Weighted Avg)", "F1-Score (Weighted Avg)"],
        "Valor": [
            100 * class_report["accuracy"],
            100 * class_report["macro avg"]["precision"],
            100 * class_report["macro avg"]["recall"],
            100 * class_report["macro avg"]["f1-score"],
            100 * class_report["weighted avg"]["precision"],
            100 * class_report["weighted avg"]["recall"],
            100 * class_report["weighted avg"]["f1-score"]
        ]
    }    
    # Imprimir en consola los resultados>
    print("Matriz de confusion -----------------------")
    print(matriz_confusion)
    print("Metricas --------------------")
    print(metricas)
    print("Reporte de clasificacion --------------------")
    print(reporte_clasif)
    print("Metricas globales -------------------------")
    print(metricas_globales)
    # Devolver el diccionario con las cuatro tablas
    return {
        "matriz_confusion": matriz_confusion,
        "metricas": metricas,
        "reporte_clasif": reporte_clasif,
        "metricas_globales": metricas_globales
    }

def get_X_y_le(tabla, columnas_referencia, colClasif):
    # Convertir el diccionario a un DataFrame de pandas
    df = pd.DataFrame(tabla)
    # Identificar columnas categóricas (tipo 'object' o 'category')
    columnas_categoricas = df[columnas_referencia].select_dtypes(include=['object', 'category']).columns.tolist()
    # Aplicar One-Hot Encoding a las columnas categóricas nominales
    if columnas_categoricas:
        transformer = ColumnTransformer(
            transformers=[
                ('onehot', OneHotEncoder(), columnas_categoricas)
            ],
            remainder='passthrough'  # Mantener las columnas no categóricas sin cambios
        )
        X = transformer.fit_transform(df[columnas_referencia])
    else:
        X = df[columnas_referencia].values
    # Codificar la columna a clasificar (colClasif) con LabelEncoder
    le = LabelEncoder()
    y = le.fit_transform(df[colClasif])
    return (X, y, le)

def Clasif_Randomforest(tabla, columnas_referencia, colClasif):
    (X, y, le) = get_X_y_le(tabla, columnas_referencia, colClasif)
    # Dividir los datos en conjunto de entrenamiento y validación (70% - 30%)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    # Crear y entrenar el modelo Random Forest
    model = RandomForestClassifier(random_state=42)
    model.fit(X_train, y_train)
    # Predecir sobre el conjunto de validación
    y_pred = model.predict(X_test)
    # Decodificar y_test y y_pred para que coincidan con los valores originales
    y_test = le.inverse_transform(y_test)
    y_pred = le.inverse_transform(y_pred)
    # Devolver la matriz de confusión y todas las métricas
    return evaluar_modelo(y_test, y_pred)
    
def Clasif_Regresionlogistica(tabla, columnas_referencia, colClasif):
    (X, y, le) = get_X_y_le(tabla, columnas_referencia, colClasif)
    # Dividir los datos en conjunto de entrenamiento y validación (70% - 30%)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    # Crear y entrenar el modelo de Regresión Logística
    model = LogisticRegression(max_iter=1000, random_state=42)  # Aumentamos max_iter para convergencia
    model.fit(X_train, y_train)
    # Predecir sobre el conjunto de validación
    y_pred = model.predict(X_test)
    # Decodificar y_test y y_pred para que coincidan con los valores originales
    y_test = le.inverse_transform(y_test)
    y_pred = le.inverse_transform(y_pred)
    # Devolver la matriz de confusión y todas las métricas
    return evaluar_modelo(y_test, y_pred)
    
def Clasif_SVM(tabla, columnas_referencia, colClasif):
    (X, y, le) = get_X_y_le(tabla, columnas_referencia, colClasif)
    # Dividir los datos en conjunto de entrenamiento y validación (70% - 30%)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    # Crear y entrenar el modelo SVM
    model = SVC(kernel='linear', random_state=42)  # Usamos un kernel lineal, pero puedes cambiarlo
    model.fit(X_train, y_train)
    # Predecir sobre el conjunto de validación
    y_pred = model.predict(X_test)
    # Decodificar y_test y y_pred para que coincidan con los valores originales
    y_test = le.inverse_transform(y_test)
    y_pred = le.inverse_transform(y_pred)
    # Devolver la matriz de confusión y todas las métricas
    return evaluar_modelo(y_test, y_pred)

def Clasif_ArbolDecision(tabla, columnas_referencia, colClasif):
    (X, y, le) = get_X_y_le(tabla, columnas_referencia, colClasif)
    # Dividir los datos en conjunto de entrenamiento y validación (70% - 30%)
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)
    # Crear y entrenar el modelo de Árbol de Decisión
    model = DecisionTreeClassifier(random_state=42)
    model.fit(X_train, y_train)
    # Predecir sobre el conjunto de validación
    y_pred = model.predict(X_test)
    # Decodificar y_test y y_pred para que coincidan con los valores originales
    y_test = le.inverse_transform(y_test)
    y_pred = le.inverse_transform(y_pred)
    # Devolver la matriz de confusión y todas las métricas
    return evaluar_modelo(y_test, y_pred)

def Clasif_default(tabla, columnas_referencia, colClasif):
    return (None, "Tipo de grafico desconocido!.")
# ------------ Funciones de grafico ------------------------------

def plotear_2dograf(plt, datosX2, datosY2, tipo_pred):
    match(tipo_pred):
        case "Generar_grafico_lineas":
            plt.plot(datosX2, datosY2, 
                   color='red', linestyle='-')
            return
        case "Generar_grafico_barras" :
            (height, msg) = extraer_num_lista(datosY2)
            if height is not None:
                plt.bar(datosX2, height=height, color='green')
            return
        case "Generar_grafico_dispersion":
            plt.scatter(datosX2, datosY2, color='blue')

def Grafico_Lineas(datosX, datosY, colX, colY, 
                   datosX2=None, datosY2=None, tipo_pred=None):
    # Convertir datosY a números
    (y, msg) = extraer_num_lista(datosY)
    if y is None:
        return (None, msg)
    # Determinar si datosX es numérico
    es_numerico = all(isinstance(x, (int, float, np.number)) and not any(pd.isna(x) for x in datosX) for x in datosX)
    # Crear figura
    plt.figure(figsize=(10, 6))
    # Configurar el gráfico de líneas
    line = plt.plot(datosX if es_numerico else range(len(datosX)), y, 
                   color='red', linestyle='-', marker='o', markersize=5)
    # Manejo inteligente del eje X
    if es_numerico:
        # Para valores numéricos
        x_vals = np.array(datosX)
        # Determinar si son enteros
        son_enteros = all(x.is_integer() for x in x_vals if isinstance(x, float))
        if son_enteros:
            # Mostrar todos los ticks si son pocos valores enteros
            if len(x_vals) <= 15:
                plt.xticks(x_vals)
            else:
                # Mostrar cada N valores si hay muchos
                step = max(1, len(x_vals) // 10)
                plt.xticks(x_vals[::step])
        else:
            # Para flotantes, mostrar ticks espaciados
            min_x, max_x = min(x_vals), max(x_vals)
            num_ticks = min(10, len(x_vals))
            plt.xticks(np.linspace(min_x, max_x, num_ticks))
        # Habilitar grid para eje X
        plt.grid(True, linestyle='--', alpha=0.7, which='both')
    else:
        # Para valores no numéricos (texto)
        if len(datosX) > 10:
            plt.xticks(range(len(datosX)), datosX, rotation=45, ha='right')
            plt.subplots_adjust(bottom=0.25)
        else:
            plt.xticks(range(len(datosX)), datosX)
        # Grid solo para eje Y
        plt.grid(True, linestyle='--', alpha=0.7, axis='y')
    # Conocer si se pasaron otros parametros
    if datosX2 is not None and datosY2 is not None and tipo_pred is not None:
        # Agregar un segundo grafico...
        plotear_2dograf(plt, datosX2, datosY2, tipo_pred)
    # Etiquetas y estilo
    plt.xlabel(colX, fontsize=12)
    plt.ylabel(colY, fontsize=12)
    # Guardar en buffer
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=120, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return (buf, "Ok")

def extraer_num_lista(datos):
    x = []
    for val in datos:
        num = extraer_numeros(val)
        if num is None:
            return (None, "Se hallo un valor en Y que no es numérico.")
        else:
            x.append(num)
    return (x, "OK")

def Grafico_Barras(datosX, datosY, colX, colY, 
                   datosX2=None, datosY2=None, tipo_pred=None):
    # Tratar de Convertir los datosY a números
    (height, msg) = extraer_num_lista(datosY)
    if height is None:
        return (None, msg)
    # Verificar si datosX es numérico (todos los elementos son int/float y no son NaN)
    es_numerico = all(
        isinstance(x, (int, float, np.number)) 
        for x in datosX) and not any(pd.isna(x) 
                                     for x in datosX)
	# Crear el gráfico de dispersión
    plt.figure(figsize=(8, 6))
    plt.bar(datosX, height=height, color='green')
    # Conocer si se pasaron otros parametros
    if datosX2 is not None and datosY2 is not None and tipo_pred is not None:
        # Agregar un segundo grafico...
        plotear_2dograf(plt, datosX2, datosY2, tipo_pred)
    # Etiquetas
    plt.xlabel(colX, fontsize=12)
    plt.ylabel(colY, fontsize=12)
    # Personalizar eje X según el tipo de dato
    if not es_numerico:
        plt.xticks(rotation=45, ha='right')  # Rotar 45° y alinear a la derecha
        plt.subplots_adjust(bottom=0.25)     # Ajustar margen inferior
	# Guardar el gráfico en un buffer de memoria
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)  # Mover el cursor al inicio del buffer
    plt.close()  # Cerrar la figura para liberar memoria
    return (buf, "Ok")

def Grafico_Scatter(datosX, datosY, colX, colY, 
                    datosX2=None, datosY2=None, tipo_pred=None):
    # Tratar de Convertir los datos a números
    (x, msgx) = extraer_num_lista(datosX)
    (y, msgy) = extraer_num_lista(datosY)
    if x is None:
        return (None, msgx)
    if y is None:
        return (None, msgy)
	# Crear el gráfico de dispersión
    plt.figure(figsize=(8, 6))
    plt.scatter(x, y, color='blue')
    # Conocer si se pasaron otros parametros
    if datosX2 is not None and datosY2 is not None and tipo_pred is not None:
        # Agregar un segundo grafico...
        plotear_2dograf(plt, datosX2, datosY2, tipo_pred)
    # Etiquetas
    plt.xlabel(colX)
    plt.ylabel(colY)
	# Guardar el gráfico en un buffer de memoria
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)  # Mover el cursor al inicio del buffer
    plt.close()  # Cerrar la figura para liberar memoria
    return (buf, "Ok")

def Grafico_Pastel(datosX, datosY, colX, colY):
    # Tratar de convertir los datosX a números
    (x, msg) = extraer_num_lista(datosX)
    if x is None:
        return (None, msg)
    # Configurar el gráfico
    fig, ax = plt.subplots(figsize=(10, 8))
    plt.subplots_adjust(right=0.7)  # Más espacio para leyenda
    # Paleta de colores
    colors = plt.cm.tab20(np.linspace(0, 1, len(x)))
    # Crear el gráfico de pastel
    wedges, texts = ax.pie(
        x,
        labels=None,
        startangle=90,
        wedgeprops={'edgecolor': 'white', 'linewidth': 1},
        colors=colors
    )
    # Calcular porcentajes
    total = sum(x)
    percentages = [f'{(value/total)*100:.1f}%' for value in x]
    # Función para ajustar dinámicamente las posiciones
    def get_text_position(angle, index):
        base_distance = 1.15  # Distancia base
        # Alternar distancia para evitar solapamiento
        variable_distance = base_distance + (0.1 if index % 2 else -0.05)
        x_pos = variable_distance * np.cos(np.deg2rad(angle))
        y_pos = variable_distance * np.sin(np.deg2rad(angle))
        return x_pos, y_pos
    # Añadir porcentajes optimizados
    for i, (p, percent) in enumerate(zip(wedges, percentages)):
        ang = (p.theta2 - p.theta1)/2. + p.theta1
        x_pos, y_pos = get_text_position(ang, i)
        ha = "left" if -90 <= ang <= 90 else "right"
        va = "center"
        # Ajuste fino para textos en posiciones verticales
        if abs(ang % 180) > 80:
            va = "bottom" if ang < 180 else "top"
        ax.text(
            x_pos, y_pos,
            percent,
            ha=ha, va=va,
            fontsize=10,
            fontweight='bold',
            bbox=dict(
                boxstyle='round,pad=0.2',
                fc='white',
                ec='none',
                alpha=0.8,
                lw=0.5
            )
        )
    # Leyenda mejorada
    legend = ax.legend(
        wedges,
        [f"{label} ({val:,})" for label, val in zip(datosY, x)],
        title=f"{colY}\nValores absolutos",
        loc="center left",
        bbox_to_anchor=(1.05, 0.5),
        frameon=True,
        framealpha=0.9,
        edgecolor='#cccccc'
    )
    # Título
    plt.title(f'Distribución de {colX}\nTotal: {total:,}', pad=20, fontsize=12)
    # Guardar imagen
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=120)
    buf.seek(0)
    plt.close()
    return (buf, "Ok")

def Grafico_Histograma(datosX, datosY, colX, colY):
    # Tratar de Convertir los datosX a números
    (x, msg) = extraer_num_lista(datosX)
    if x is None:
        return (None, msg)
	# Crear el gráfico de pastel
    plt.figure(figsize=(8, 6))
    plt.hist(x, bins=10, color='blue')
    plt.xlabel(colX)
    plt.ylabel("Frecuencia")
	# Guardar el gráfico en un buffer de memoria
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)  # Mover el cursor al inicio del buffer
    plt.close()  # Cerrar la figura para liberar memoria
    return (buf, "Ok")

# ----------------------------------
import dask.dataframe as dd
import pandas as pd
from docx import Document
import pdfplumber
import json
import os

#------- Funciones auxiliares y principal para la extracción de tablas en los documetos -----------------

def es_tabla_valida(df):
    """Verifica si un DataFrame es válido (no está vacío y tiene al menos una fila y columna)."""
    return not df.empty and df.shape[0] > 0 and df.shape[1] > 0

def convertir_a_json(df):
    """Convierte un DataFrame a un diccionario donde las claves son los nombres de las columnas
    y los valores son listas de elementos de cada columna."""
    if es_tabla_valida(df):
        return {col: df[col].tolist() for col in df.columns}  # Formato deseado
    return None

def es_fila_valida(fila):
    """Verifica si una fila contiene al menos un valor no NaN."""
    return any(cell is not None and not (isinstance(cell, float) and np.isnan(cell)) for cell in fila)

def detectar_filas_encabezados(data):
    """Detecta automáticamente el número de filas de encabezados."""
    if len(data) < 2:
        return 1
    def es_fila_de_datos(fila):
        count_str = 0
        count_float_dt = 0
        count_nan_empty = 0
        for cell in fila:
            # Caso 1: Manejar NaN/None explícitamente (pd.NA, np.nan, None)
            if pd.isna(cell) or cell == "":
                count_nan_empty += 1
                continue
            # Caso 2: Verificar números (enteros/floats, incluyendo cadenas numéricas)
            try:
                if not extraer_numeros(cell) is None:
                    count_float_dt += 1
                    continue
                #return True
            except (ValueError, TypeError):
                pass
            # Caso 3: Verificar fechas (opcional)
            if isinstance(cell, str):
                try:
                    pd.to_datetime(cell)
                    count_float_dt += 1
                    continue
                    #return True
                except:
                    # es string o categoria
                    count_str += 1
                    pass
        sum_all = (count_float_dt+count_str+count_nan_empty)
        #if count_float_dt > count_str and count_float_dt > count_nan_empty:
        if count_float_dt > 1:
            return (True, count_float_dt / sum_all)
        return (False, count_str / sum_all)
    # Buscar la primera fila que cumpla con es_fila_de_datos
    i = 0
    fila_encabezados = None
    while i < len(data):
        fila = data[i]
        (resp, pc) = es_fila_de_datos(fila)
        if resp:
            if not fila_encabezados is None:
                return fila_encabezados+1
            # Comprobar si la siguiente sigue siendo fila de datos...
            i += 1
            if i < len(data):
                 fila = data[i]
                 (resp, pc) = es_fila_de_datos(fila)
                 if resp:
                     return i  # Devuelve el índice de la última fila de encabezado
                 else:
                     fila_encabezados = i
        else:
            if pc == 1.0:
                return i+1
            else:
                fila_encabezados = i
        i += 1
    return 0  # Si no encuentra fila de datos, asume que todas son encabezados

def detectar_tabla_principal(data):
    """Detecta la tabla más grande en una hoja de Excel, asegurándose de que no haya filas sin datos
    y que las columnas tengan datos válidos. Devuelve solo una tabla por hoja."""
    bloques = []
    inicio_fila = 0
    while inicio_fila < len(data):
        # Ignorar filas con solo NaN
        if not es_fila_valida(data[inicio_fila]):
            inicio_fila += 1
            continue
        # Buscar la primera fila que parece un encabezado (contiene texto)
        if any(cell and isinstance(cell, str) for cell in data[inicio_fila]):
            fin_fila = inicio_fila + 1
            # Buscar el final del bloque tabular
            while fin_fila < len(data) and es_fila_valida(data[fin_fila]):
                fin_fila += 1
            # Verificar que todas las filas dentro del bloque tengan datos válidos
            bloque_valido = True
            for fila in data[inicio_fila:fin_fila]:
                if not es_fila_valida(fila):
                    bloque_valido = False
                    break
            if bloque_valido and fin_fila - inicio_fila > 1:
                # Detectar desde qué columna y hasta qué columna hay datos válidos
                inicio_col = 0
                fin_col = len(data[0])  # Asumimos que todas las filas tienen el mismo número de columnas
                # Encontrar la primera columna con datos válidos
                for j in range(len(data[0])):
                    if any(es_fila_valida([data[i][j]]) for i in range(inicio_fila, fin_fila)):
                        inicio_col = j
                        break
                # Encontrar la última columna con datos válidos
                for j in range(len(data[0]) - 1, -1, -1):
                    if any(es_fila_valida([data[i][j]]) for i in range(inicio_fila, fin_fila)):
                        fin_col = j + 1
                        break
                # Añadir el bloque recortado (filas y columnas válidas)
                bloques.append((inicio_fila, fin_fila, inicio_col, fin_col))
            inicio_fila = fin_fila
        else:
            inicio_fila += 1
    # Seleccionar el bloque más grande (el que tenga más filas)
    if bloques:
        bloques.sort(key=lambda x: x[1] - x[0], reverse=True)  # Ordenar por tamaño (filas)
        return bloques[0]  # Devolver el bloque más grande
    return None  # No se encontró ningún bloque válido

def procesar_encabezados(data):
    """Procesa filas de encabezados múltiples y celdas vacías."""
    # Rellenar celdas vacías en los encabezados
    for i in range(len(data) - 1):
        for j in range(len(data[i])):
            if not data[i][j]:
                data[i][j] = data[i - 1][j] if i > 0 else f"Columna_{j}"
    # Combinar filas de encabezados
    headers = []
    for j in range(len(data[0])):
        header = "_".join([str(data[i][j]) for i in range(len(data)) if data[i][j]])
        headers.append(header)
    # Hacer nombres de columnas únicos
    unique_headers = []
    seen = {}
    for header in headers:
        if header in seen:
            seen[header] += 1
            unique_headers.append(f"{header}_{seen[header]}")
        else:
            seen[header] = 0
            unique_headers.append(header)
    return unique_headers

def leer_archivo_grande(tmp_doc, es_excel=True):
    """Lee archivos grandes usando Dask."""
    if es_excel:
        df_dict = dd.read_excel(tmp_doc, sheet_name=None, header=None)
        return {sheet: df.compute() for sheet, df in df_dict.items()}
    else:
        df_sample = pd.read_csv(tmp_doc, nrows=100, header=None)
        dtypes = {col: 'float64' if pd.to_numeric(df_sample[col], errors='coerce').notna().all() else 'object'
                  for col in df_sample.columns}
        return dd.read_csv(tmp_doc, dtype=dtypes).compute()

def dar_tablas(tmp_doc, ext_doc):
    # funciones auxiliares clean_cell y homogenize_dataframe
    def clean_cell(value):
        """
        Función para estandarizar el valor de una celda según los requisitos:
        - Textos (incluyendo fechas, categorías) se mantienen como str
        - "-" (con posibles espacios) se convierte a 0
        - Vacíos se mantienen como NaN
        - Números (con comas/puntos) se convierten a float/int
        """
        # Caso 1: Valor es NaN o vacío
        if pd.isna(value) or str(value).strip() == '':
            return np.nan
        str_val = str(value).strip()
        # Caso 2: Es un guión "-"
        if str_val == '-':
            return 0
        # Caso 3: Es un número (puede contener comas/puntos como separadores)
        if str_val.replace(',', '').replace('.', '').isdigit():
            try:
                # Primero quitamos comas (separadores de miles)
                clean_num = str_val.replace(',', '')
                # Convertimos a float (manejará correctamente el punto decimal)
                num_val = float(clean_num)
                # Si es entero, lo devolvemos como int
                return int(num_val) if num_val.is_integer() else num_val
            except:
                return str_val  # Si falla la conversión, mantener original
        # Caso 4: Es texto (lo mantenemos igual)
        return str_val

    def homogenize_dataframe(df):
        """
        Aplica la limpieza celda por celda a un DataFrame completo
        """
        return df.map(clean_cell)
    
    tablas = []
    try:
        if ext_doc in ['.doc', '.docx']:  # ============== OFFICE WORD =====================
            doc = Document(tmp_doc)
            for tabla in doc.tables:
                data = [[cell.text.strip() if cell.text else "" for cell in row.cells] for row in tabla.rows]
                num_encabezados = detectar_filas_encabezados(data)
                if num_encabezados:
                    headers = procesar_encabezados(data[:num_encabezados])  # Procesar las filas de encabezados detectadas
                    df = pd.DataFrame(data[num_encabezados:], columns=headers)  # Resto de filas como datos
                    tablas.append(convertir_a_json(df))
        elif ext_doc in ['.xls', '.xlsx', '.xlsm', '.xlsb', '.odf', '.ods', '.odt']:  # ============== OFFICE EXCEL =====================
            if os.path.getsize(tmp_doc) > 100 * 1024 * 1024:  # Archivo grande
                df_dict = leer_archivo_grande(tmp_doc, es_excel=True)
            else:
                df_dict = pd.read_excel(
                    tmp_doc,
                    sheet_name=None,
                    header=None,
                    engine='openpyxl',  # Usar openpyxl para mejor compatibilidad
                    #dtype=str,  # Forzar todo a string para evitar problemas
                    na_values=['', 'NA', 'N/A'],  # Valores a considerar como NaN
                    keep_default_na=False  # Evitar convertir strings como 'NULL' a NaN
                )
            for sheet, df in df_dict.items():
                if not df.empty:
                    # Convertir el DataFrame a una lista de listas
                    data = df.values.tolist()
                    # Detectar la tabla principal en la hoja
                    bloque = detectar_tabla_principal(data)
                    if bloque:
                        inicio_fila, fin_fila, inicio_col, fin_col = bloque
                        bloque_data = [fila[inicio_col:fin_col] for fila in data[inicio_fila:fin_fila]]
                        num_encabezados = detectar_filas_encabezados(bloque_data)
                        if num_encabezados:
                            headers = procesar_encabezados(bloque_data[:num_encabezados])  # Procesar las filas de encabezados detectadas
                            df_bloque = pd.DataFrame(bloque_data[num_encabezados:], columns=headers)  # Resto de filas como datos
                            df_bloque = df_bloque.fillna("")
                            tablas.append(convertir_a_json(df_bloque))
        elif ext_doc == '.pdf':  # ============== PDF =====================
            with pdfplumber.open(tmp_doc) as pdf:
                for page in pdf.pages:
                    for tabla in page.extract_tables():
                        if tabla:  # Verificar si la tabla no está vacía
                            num_encabezados = detectar_filas_encabezados(tabla)
                            if num_encabezados:
                                headers = procesar_encabezados(tabla[:num_encabezados])  # Procesar las filas de encabezados detectadas
                                df = pd.DataFrame(tabla[num_encabezados:], columns=headers)  # Resto de filas como datos
                                tablas.append(convertir_a_json(df))
        elif ext_doc == '.csv':  # ============== CSV =====================
            if os.path.getsize(tmp_doc) > 100 * 1024 * 1024:  # Archivo grande
                df = leer_archivo_grande(tmp_doc, es_excel=False)
            else:
                try:
                    df = pd.read_csv(   tmp_doc, 
                                        header=None, 
                                        dtype=str,
                                        sep=None,  # Auto-detecta separador
                                        engine='python',
                                        na_values=[""],
                                        keep_default_na=False,
                                        on_bad_lines='warn'
                    )
                except UnicodeDecodeError:
                    df = pd.read_csv(   tmp_doc, 
                                        header=None, 
                                        dtype=str,
                                        sep=None,  # Auto-detecta separador
                                        engine='python',
                                        na_values=[""],
                                        keep_default_na=False,
                                        encoding='latin-1',
                                        on_bad_lines='warn'
                    )
                except pd.errors.ParserError:
                    df = pd.read_csv(   tmp_doc, 
                                        header=None, 
                                        dtype=str,
                                        sep=';',
                                        engine='python',
                                        na_values=[""],
                                        keep_default_na=False,
                                        encoding='latin-1',
                                        on_bad_lines='warn'
                    )
            if not df.empty:
                # Aplicar limpieza homogénea
                df_clean = homogenize_dataframe(df)
                # Convertir el DataFrame a una lista de listas
                data = df_clean.values.tolist()
                # Detectar la tabla principal en el archivo CSV
                bloque = detectar_tabla_principal(data)
                if bloque:
                    inicio_fila, fin_fila, inicio_col, fin_col = bloque
                    bloque_data = [fila[inicio_col:fin_col] for fila in data[inicio_fila:fin_fila]]
                    num_encabezados = detectar_filas_encabezados(bloque_data)
                    if num_encabezados:
                        headers = procesar_encabezados(bloque_data[:num_encabezados])  # Procesar las filas de encabezados detectadas
                        df_bloque = pd.DataFrame(bloque_data[num_encabezados:], columns=headers)  # Resto de filas como datos
                        df_bloque = df_bloque.fillna("")
                        tablas.append(convertir_a_json(df_bloque))
        elif ext_doc == '.json':  # ============== JSON =====================
            with open(tmp_doc, 'r', encoding='utf-8') as file:
                tablas = json.load(file)
        else:
            return json.dumps({"error": "Extensión no reconocida."})
        # Filtrar tablas no válidas (None)
        tablas = [tabla for tabla in tablas if tabla is not None]
        return json.dumps(tablas)
    except Exception as e:
        error_msg = {"error": f"Error al procesar el archivo: {str(e)}"}
        return json.dumps(error_msg, ensure_ascii=False)

# ================== P U N T O S   D E   E N T R A D A   A L   S E R V I D O R   F L A S K =================

@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"error": "El archivo es demasiado grande. El límite es 20 GB."}), 413

@app.route("/")
def init():
    if session.get('first_time') is None:
        reset_global_variables()
        session['first_time'] = False
        gvar = get_global_variables()
        textos = [
            "Descubra el poder de los datos con nuestro sistema de BIG DATA Documents. "
            "Transforme información en conocimiento y tome decisiones inteligentes "
            "mediante Informes KPI (Key Performance Indicators). "
            "Analice, visualice y optimice el rendimiento de su organización "
            "de manera eficiente y efectiva. ¡Convierta los datos en su mejor aliado!",

            "En un mundo impulsado por datos, la clave del éxito está en transformar información en acción. "
            "Con BIG DATA Documents, potencie su organización mediante Informes KPI (Key Performance Indicators) "
            "que revelan insights poderosos. Analice tendencias, optimice procesos y tome decisiones estratégicas "
            "con precisión. ¡Convierta los datos en su ventaja competitiva y lleve su negocio al siguiente nivel!"
        ]
        texto_inicial = random.choice(textos)  # Elige un texto aleatorio
        return render_template("principal.html", Start=True, texto_inicial=Markup(texto_inicial), **gvar)
    else:
        return inicio()

def inicio():
    gvar = get_global_variables()
    textos = [
        "Desbloquee el potencial oculto en sus datos con nuestro sistema BIG DATA Documents. "
        "Convierta la información en conocimiento valioso y tome decisiones más inteligentes y "
        "fundamentadas utilizando Informes KPI (Indicadores Clave de Rendimiento). Explore, "
        "analice y visualice el rendimiento de su organización de manera clara y eficiente. "
        "¡Haga que los datos trabajen para usted y conviértalos en su herramienta más poderosa para el éxito!",

        "En un mundo donde los datos son el motor del progreso, la clave para destacar es convertir la "
        "información en acciones concretas. Con BIG DATA Documents, impulse su organización hacia adelante"
        " utilizando Informes KPI que desvelan insights transformadores. Detecte tendencias, mejore sus "
        "procesos y tome decisiones estratégicas con confianza. ¡Aproveche el poder de los datos para ganar "
        "ventaja y llevar su negocio a nuevas alturas!"
    ]
    texto_inicial = random.choice(textos)  # Elige un texto aleatorio
    return render_template('inicio.html', Start=False, texto_inicial=Markup(texto_inicial), **gvar)

@app.route("/login", methods=['GET','POST'])
def autenticar():
    global conn
    gvar = get_global_variables()
    if request.method == 'POST':
        #Verificar si el usuario y contrase~na introducidas estan registradoa...
        username = request.form['username']
        password = request.form['password']
        query = text("SELECT * FROM usuarios WHERE username='"+username+"' AND password='"+password+"'")
        #verificar primero si la conexion a la base de datos esta valida, sino restaurar
        try:
            result = conn.execute(query).fetchone()
        except SQLAlchemyError as e:
            conn = mydb.connect()
            result = conn.execute(query).fetchone()
        if result!=None:
            if result.aprobado:
                tipo_usuario = result.rol
                if result.url_foto != "" and result.url_foto is not None:
                    url_foto = 'fotos_usuarios/'+result.url_foto
                else:
                    url_foto = 'img/anonimo_black.png'
                # Se toma por defecto el primer chatbot con codigo 1 (spaCy)
                set_global_variables(username, password, url_foto, tipo_usuario)
                session.permanent = True  # Marcar la sesión como permanente
                return inicio()
            else:
                return render_template('mensajes_de_error.html',encabezado=Markup('Usuario aún no aprobado'),
                mensaje=Markup('Usuario <strong>'+result.nom_ape+'</strong>: Usted ya está registrado, pero debe esperar aprobación por los administradores.'), **gvar)    
        else:
            return render_template('mensajes_de_error.html',encabezado='Usuario no existente',
                mensaje=Markup('Ha introducido un nombre de usuario o contraseña incorrectos.'), **gvar) 
    else: #GET
        #Mostrar el formulario de acceder al sistema
        return render_template('login.html',Start=False, **gvar)

# Salir de la cuenta autenticada
@app.route("/logout")
def logout():
    reset_global_variables()
    return init()

# Registrarse como usuario nuevo
@app.route("/registrarse", methods=['GET','POST'])
def registrarse():
    global conn
    gvar = get_global_variables()
    if request.method == 'POST':
        nombres = request.form["nombres"]
        apellidos = request.form["apellidos"]
        email = request.form["email"]
        username = request.form["username"]
        password = request.form["password"]
        
        # Chequear que no haya un usuario con el mismo nombres/apellidos o username
        query = text("SELECT * FROM usuarios WHERE nom_ape='"+nombres+" "+apellidos+"' OR username='"+username+"'")
        try:
            result = conn.execute(query).fetchall()
        except SQLAlchemyError as e:
            conn = mydb.connect()
            result = conn.execute(query).fetchall()
        if len(result):
            return render_template('mensajes_de_error.html',encabezado=Markup('Error de nuevo registro'),
                mensaje=Markup('Debe especificar nombres/apellidos o usuario (<em>username</em>) distintos.'), **gvar)
            
        archivo_foto = request.files['url_foto']
        usuarios = metadata.tables['usuarios']
        oper = usuarios.insert().values(nom_ape=nombres+' '+apellidos,email=email,
                rol='gestor_informes',username=username,password=password)
        try:
            result = conn.execute(oper)
        except SQLAlchemyError as e:
            conn = mydb.connect()
            result = conn.execute(oper)
        conn.commit()
        #recuperando el id para el nuevo usuario o el existente...
        id_usuario = result.lastrowid
        #Guardando al servidor la foto del usuario...
        if archivo_foto.filename != '':
            #componemos el nombre id.[ext foto] y lo archivamos en el servidor...
            ext_foto = archivo_foto.filename.rsplit('.', 1)[1]
            url_foto = str(id_usuario)+'.'+ext_foto
            definitive_url_foto = 'static/fotos_usuarios/' + url_foto
            archivo_foto.save(definitive_url_foto)
            #ahora, actualizamos el campo url_foto para ese usuario..
            upd = usuarios.update().where(usuarios.c.idUsuario==id_usuario).values(url_foto=url_foto)
            conn.execute(upd)
            conn.commit()
        return inicio()
    else: # GET
        return render_template('registrarse.html', Start=False, **gvar)

@app.route("/modelos")
@app.route("/modelos/<valor_sel>")
def modelos(valor_sel = "-1"):
    global conn
    gvar = get_global_variables()
    if gvar["username"] == "":
        return autenticar()
    # Obtener datos usuario de la sesion
    query = text('SELECT * FROM usuarios WHERE username="'+gvar['username']+'" AND password="'+gvar['password']+'"')
    try:
        user = conn.execute(query).fetchone()
    except SQLAlchemyError as e:
        conn = mydb.connect()
        user = conn.execute(query).fetchone()
    default_options = {"-1":"Ninguno", "0":"Nuevo modelo"}
    lista_modelos = ''
    for key, value in default_options.items():
        if key == valor_sel:
            sel = 'selected="selected"'
        else:
            sel = ''
        lista_modelos += '<option value="'+key+'" '+sel+'>'+value+'</option>'
    # buscar los modelos creados por el usuario actual para actualizar la lista de modelos...    
    query = text('SELECT * FROM modelos WHERE idUsuario='+str(user.idUsuario))
    res = conn.execute(query).fetchall()
    if not res is None:
        for row in res:
            lista_modelos += '<option value="'+str(row.idUsuario)+'_'+str(row.idModelo)+'.json">'+row.nombre+'</option>'
    return render_template('modelos.html', lista_modelos=Markup(lista_modelos), **gvar)

@app.route("/leer_archivo_json", methods=['GET','POST'])
def leer_archivo_json():
    if request.method == 'POST':
        # Leemos el archivo json y lo enviamos al front-end...
        archivo_json = request.form["archivo_json"]
        path_name = 'static/json/'+archivo_json
        with open(path_name, 'r', encoding='utf-8') as file:
            data = json.load(file)
        return data

@app.route("/salvar_modelo", methods=['POST'])
def salvar_modelo():
    global conn
    gvar = get_global_variables()
    # Obtener datos usuario de la sesion
    query = text('SELECT * FROM usuarios WHERE username="'+gvar['username']+'" AND password="'+gvar['password']+'"')
    try:
        user = conn.execute(query).fetchone()
    except SQLAlchemyError as e:
        conn = mydb.connect()
        user = conn.execute(query).fetchone()
    if request.method == 'POST':
        listas_json = request.form["listas_json"]
        nombre_modelo = request.form["nombre_modelo"]
        modelos = metadata.tables['modelos']
        # Chequear si ya existe un modelo con el nombre pasado para con el usuario actual...
        query = text('SELECT * FROM modelos WHERE idUsuario='+str(user.idUsuario)+' AND nombre="'+nombre_modelo+'"')
        res = conn.execute(query).fetchone()
        if not res is None:
            id_modelo = res.idModelo
        else:
            # Guardar a la tabla modelos...
            ins = modelos.insert().values(idUsuario=user.idUsuario,nombre=nombre_modelo)
            res = conn.execute(ins)
            conn.commit()
            id_modelo = res.lastrowid
        if id_modelo:
            # Vamos primero a purgar la lista, en el elemento workflow_json...
            listas = json.loads(listas_json)
            workflow_json = listas[0]
            conections = listas[1]
            for obj in workflow_json.copy():
                conect = next(
                    (item for item in conections if item["from"]==obj["id"] or item["to"]==obj["id"]), 
                    None  # Valor por defecto si no se encuentra
                )
                if conect is None:
                    # Eliminar objeto que no tiene conecciones
                    workflow_json.remove(obj)
            #volvemos a jsonniciar ...
            listas = [workflow_json, conections]
            listas_json = json.dumps(listas)
            # Guardamos el nuevo archivo json...
            archivo_json = 'static/json/'+str(user.idUsuario)+'_'+str(id_modelo)+'.json'
            try:
                with open(archivo_json,'wt') as file:
                    json.dump(listas_json, file)
                # actualizar en tabla modelos...
                upd = modelos.update().where(modelos.c.idModelo==id_modelo).values(archivo_json=archivo_json)
                res = conn.execute(upd)
                conn.commit()
                return json.dumps({"Status" : "OK"})
            except Exception as e:
                # eliminar fila guardada
                query = text("DELETE FROM modelos WHERE idModelo="+str(id_modelo))
                conn.execute(query)
                conn.commit()
                return json.dumps({"Status" : "Error", "Message" : f"No se pudo guardar el archivo JSON: {str(e)}"})
        else:
            return json.dumps({"Status" : "Error", "Message" : "No se pudo archivar el modelo!."})

@app.route("/leer_doc_desde_pc", methods=['GET','POST'])
def leer_doc_desde_pc():
    global idx_tmp
    gvar = get_global_variables()
    if request.method == 'POST':
        # Leemos el documento cargado y lo salvamos temporalmente en el servidor...
        documento = request.files["documento"]
        ext_doc = os.path.splitext(documento.filename)[1].lower()
        idx_tmp += 1
        tmp_doc = 'static/tmp/tmp' + str(idx_tmp) + ext_doc
        documento.save(tmp_doc)
        tablas = dar_tablas(tmp_doc, ext_doc)
        return Response(tablas, mimetype='application/json')
    else:
        return render_template('mensajes_de_error.html',encabezado=Markup('Error con uso de los grupos tabulados'),
                mensaje=Markup('Se ha enviado un metodo de respuesta que no es POST.'), **gvar)

@app.route("/leer_doc_desde_url", methods=['POST'])
def leer_doc_desde_url():
    global idx_tmp
    gvar = get_global_variables()
    if request.method == 'POST':
        url = request.form["url"]
        try:
            # Hacer una solicitud GET a la URL
            respuesta = requests.get(url)            
            # Verificar si la solicitud fue exitosa (código 200)
            if respuesta.status_code == 200:
                ext_doc = os.path.splitext(url)[1].lower()
                idx_tmp += 1
                tmp_doc = 'static/tmp/tmp' + str(idx_tmp) + ext_doc
                # Escribir el contenido en un archivo local
                with open(tmp_doc, 'wb') as archivo:
                    archivo.write(respuesta.content)
                return dar_tablas(tmp_doc, ext_doc)
            else:
                print(f"Error al descargar el archivo. Código de estado: {respuesta.status_code}")
                return 0
        except requests.exceptions.RequestException as e:
            print(f"Error en la solicitud: {e}")
            return 0
    else:
        return render_template('mensajes_de_error.html',encabezado=Markup('Error con uso de los grupos tabulados'),
                mensaje=Markup('Se ha enviado un metodo de respuesta que no es POST.'), **gvar)

@app.route("/tabla_operaciones/<operacion>", methods=['POST'])
def tabla_operaciones(operacion):
    try:
        # Recibir los datos JSON enviados desde el cliente
        datos = request.get_json()
        # Conmutar cual operacion fue la enviada sobre la tabla
        match(operacion):
            case "Quitar_duplicados":
                tabla = datos['tabla']
                columnas_referencia = datos['columnasReferencia']
                tablaModificada = Quitar_duplicados(tabla, columnas_referencia)
                # Convertir el DataFrame de nuevo a un diccionario
                tablaModificada = convertir_a_json(tablaModificada) # tablaModificada.to_json(orient='split')
                # Devolver la tabla modificada al cliente
                print(Procesos_Label2Descrip[operacion]," ============= ")
                return jsonify({
                    "mensaje": "Duplicados eliminados correctamente",
                    "tablaModificada": tablaModificada
                }), 200
            case "Recuperar_num_corruptos":
                tabla = datos['tabla']
                columnas_referencia = datos['columnasReferencia']
                tablaModificada = RecuperarNumCorruptos(tabla, columnas_referencia)
                # Convertir el DataFrame de nuevo a un diccionario
                tablaModificada = convertir_a_json(tablaModificada) #tablaModificada.to_json(orient='split')
                # Devolver la tabla modificada al cliente
                print(Procesos_Label2Descrip[operacion]," ============= ")
                print(tablaModificada)
                return jsonify({
                    "mensaje": "Recuperacion de numeros realizado.",
                    "tablaModificada": tablaModificada
                }), 200
            case "Prediccion":
                tabla = datos["tabla"]                
                colX = datos["columnasXY"][0]
                colY = datos["columnasXY"][1]
                epocas = int(datos["epocas"])
                proceso = datos["proceso"]
                print(Procesos_Label2Descrip[proceso]," ============= ")
                match(proceso):
                    case "Prediccion_lineal":
                        # Aplicar prediccion lineal...
                        df = Prediccion_lineal(tabla[colX], tabla[colY], colX, colY, epocas)
                        # Convertir a JSON
                        json_resultado = convertir_a_json(df)
                        # Devolver los resultados al cliente
                        return jsonify({
                            "mensaje": "Prediccion lineal realizada",
                            "predicciones": json_resultado
                        }), 200
                    case "Prediccion_nolineal":
                        # Aplicar prediccion lineal...
                        df = Prediccion_nolineal(tabla[colX], tabla[colY], colX, colY, epocas)
                        # Convertir a JSON
                        json_resultado = convertir_a_json(df)
                        # Devolver los resultados al cliente
                        return jsonify({
                            "mensaje": "Prediccion no lineal realizada",
                            "predicciones": json_resultado
                        }), 200
                    case "Prediccion_SeriesTemporales":
                        # Preparar datos para Prophet
                        df = pd.DataFrame(tabla)
                        predicciones = Prediccion_SeriesTemporales(df, colX, colY, epocas)
                        # Devolver los resultados al cliente
                        predicciones = convertir_a_json( predicciones )
                        return jsonify({
                            "mensaje": "Prediccion por Series Temporales realizada",
                            "predicciones": predicciones
                        }), 200
            case "Estadistica_descriptiva":
                tabla = datos['tabla']
                columnas_referencia = datos['columnasReferencia']
                tablaModificada = EstadisticaDescriptiva(tabla, columnas_referencia)
                # Convertir el DataFrame de nuevo a un diccionario
                tablaModificada = convertir_a_json(tablaModificada)
                # Devolver la tabla modificada al cliente
                print(Procesos_Label2Descrip[operacion]," ============= ")
                return jsonify({
                    "mensaje": "Estadistica Descriptiva realizada.",
                    "tablaModificada": tablaModificada
                }), 200
            case "Clasificacion_predictiva":
                tabla = datos['tabla']
                columnas_referencia = datos['columnasReferencia']
                colClasif = datos['colClasif']
                proceso = datos['proceso']
                print(Procesos_Label2Descrip[proceso]," ============= ")
                match(proceso):
                    case "Clasif_Arbol_decision":
                        # Aplicar Arbol de decision...
                        resultados = Clasif_ArbolDecision(tabla, columnas_referencia, colClasif)
                        # Devolver los resultados al cliente
                        return jsonify({
                            "mensaje": "Clasificación por Árbol de Decisión existosa",
                            "resultados": resultados
                        }), 200
                    case "Clasif_Random_forest":
                        # Aplicar Arbol de decision...
                        resultados = Clasif_Randomforest(tabla, columnas_referencia, colClasif)
                        # Devolver los resultados al cliente
                        return jsonify({
                            "mensaje": "Clasificación por Bosque Aleatorio (Random Forest) existosa",
                            "resultados": resultados
                        }), 200
                    case "Clasif_Regresion_logistica":
                        # Aplicar Arbol de decision...
                        resultados = Clasif_Regresionlogistica(tabla, columnas_referencia, colClasif)
                        # Devolver los resultados al cliente
                        return jsonify({
                            "mensaje": "Clasificación por Regresión Logística existosa",
                            "resultados": resultados
                        }), 200
                    case "Clasif_SVM":
                        # Aplicar Arbol de decision...
                        resultados = Clasif_SVM(tabla, columnas_referencia, colClasif)
                        # Devolver los resultados al cliente
                        return jsonify({
                            "mensaje": "Clasificación por Máquina de Vectores de Soporte existosa",
                            "resultados": resultados
                        }), 200
            case _:
                return jsonify({
                            "error": "Operacion no reconocida",
                        }), 400
    except Exception as e:
        return jsonify({"error": str(e)}), 400

# Tabla LookUp para saltar y llamar a la funcion grafico correspondiente
accTipoGraf = {
    "Generar_grafico_lineas" : Grafico_Lineas,
    "Generar_grafico_barras" : Grafico_Barras,
    "Generar_grafico_dispersion" : Grafico_Scatter,
    "Generar_grafico_pastel" : Grafico_Pastel,
    "Generar_grafico_histograma" : Grafico_Histograma
}
def accion_grafdefault(tabla, columnas):
    return (None,"Hay un Objeto en el modelo que no identificado.")

@app.route("/generar_grafico", methods=['POST'])
def generar_grafico():
    try:
        # Recibir los datos JSON enviados desde el cliente
        datos = request.get_json()
        tabla = datos["tabla"]
        columnas = datos["columnas"]
        colX = columnas[0]
        colY = columnas[1]
        # Extraer los datos de las columnas X e Y
        datosX = tabla[colX]
        datosY = tabla[colY]
        tipo = datos["type"]
        fun_graf = accTipoGraf.get(tipo, accion_grafdefault)
        # Antes de seguir, averiguemos si se envio' otros X y Y para 
        # que se grafiquen en el mismo eje y con cuales tipo de grafico...
        if "type_pred" in datos.keys():
            tabla_pred = datos["tabla_pred"]
            datosX2 = tabla_pred[colX]
            datosY2 = tabla_pred[colY]
            tipo_pred = datos["type_pred"]
            # Graficado doble
            (buf, msg) = fun_graf(datosX, datosY, colX, colY, 
                                  datosX2, datosY2, tipo_pred)
        else:
            # Graficado sencillo
            (buf, msg) = fun_graf(datosX, datosY, colX, colY)
        if buf is None:
            return jsonify({
                "error": msg,
            }), 400
        else:
            # Devolver la imagen como un blob
            return send_file(buf, mimetype='image/png')
    except Exception as e:
        return jsonify({"error": str(e)}), 400

import pandas as pd

def agrupar_filas(data):
    """
    Función completa para agrupamiento que maneja:
    - count_categories: Conteo de cada valor categórico por grupo
    - count: Conteo total de registros por grupo
    - count_distinct: Conteo de valores únicos por grupo
    - Operaciones numéricas: sum, mean, max, min
    
    Parámetros:
    data : dict
        Debe contener:
        - table: Diccionario con los datos
        - groupBy: Lista de columnas para agrupar
        - operations: Lista de diccionarios con 'column' y 'operation'
    
    Retorna:
    dict : Resultado en formato JSON compatible con convertir_a_json
    """
    # Convertir a DataFrame
    df = pd.DataFrame(data['table'])
    
    # Validaciones iniciales
    if not isinstance(data.get('groupBy', []), list):
        raise ValueError("groupBy debe ser una lista de columnas")
    if not isinstance(data.get('operations', []), list):
        raise ValueError("operations debe ser una lista de operaciones")
    
    # Verificar existencia de columnas
    all_columns = set(df.columns)
    for col in data['groupBy']:
        if col not in all_columns:
            raise ValueError(f"Columna de agrupación '{col}' no existe")
    
    # Preparar resultados
    results = []
    
    for op in data['operations']:
        col = op['column']
        operation = op['operation'].lower()
        
        if col not in all_columns:
            raise ValueError(f"Columna de operación '{col}' no existe")
        
        # Operación: Conteo por categorías (APROBADA: X, EN REVISION: Y)
        if operation == 'count_categories':
            # Agrupar por provincia + estado y contar ocurrencias
            count_df = (df.groupby(data['groupBy'] + [col])
                       .size()
                       .unstack(fill_value=0)
                       .reset_index())
            results.append(count_df)
        
        # Operación: Conteo total de registros
        elif operation == 'count':
            count_df = df.groupby(data['groupBy']).size().reset_index(name=f'{col}_count')
            results.append(count_df)
        
        # Operación: Conteo de valores distintos
        elif operation == 'count_distinct':
            distinct_df = (df.groupby(data['groupBy'])[col]
                          .nunique()
                          .reset_index(name=f'{col}_distinct'))
            results.append(distinct_df)
        
        # Operaciones numéricas
        elif operation in ['sum', 'mean', 'max', 'min']:
            try:
                df[col] = pd.to_numeric(df[col], errors='raise')
                op_df = (df.groupby(data['groupBy'])[col]
                         .agg(operation)
                         .reset_index(name=f'{col}_{operation}'))
                results.append(op_df)
            except ValueError:
                raise ValueError(f"Columna '{col}' no es numérica para {operation}")
        
        else:
            raise ValueError(f"Operación no soportada: '{operation}'")
    
    # Combinar todos los resultados
    if not results:
        raise ValueError("No se generaron resultados válidos")
    
    final_result = results[0]
    for df in results[1:]:
        final_result = final_result.merge(df, on=data['groupBy'], how='outer')
    
    # Convertir a formato JSON
    return convertir_a_json(final_result)

@app.route("/groupby", methods=['POST'])
def groupby():
    data = request.get_json()
    nueva_tabla = agrupar_filas( data )
    nueva_tabla = json.dumps( nueva_tabla )
    return Response(nueva_tabla, mimetype='application/json')

def ordenar_filas(data):
    df = pd.DataFrame(data['table'])
    # Parámetros de ordenamiento
    sort_by = data['sort_by']  # Ej: "Edad" o ["Ciudad", "Edad"]
    ascending = data.get('ascending', True)  # Default: ascendente
    # Ordenar
    df_sorted = df.sort_values(by=sort_by, ascending=ascending)
    tabla_json = convertir_a_json( df_sorted )
    return tabla_json

@app.route("/ordenar_tabla", methods=['POST'])
def ordenar_tabla():
    data = request.get_json()
    tabla_json = ordenar_filas( data )    
    nueva_tabla = json.dumps( tabla_json )
    return Response(nueva_tabla, mimetype='application/json')

def eliminar_filas(data):
    df = pd.DataFrame(data['table'])
    column = data['column']
    operator_str = data['operator']  # Renombrado para evitar conflicto con el módulo operator
    value = data['value']
    case_sensitive = data.get('case_sensitive', False)  # Valor por defecto False si no viene
    # Intentar convertir el valor al tipo de dato de la columna
    try:
        if pd.api.types.is_numeric_dtype(df[column]):
            value = float(value) if '.' in value else int(value)
        elif pd.api.types.is_datetime64_any_dtype(df[column]):
            value = pd.to_datetime(value)
    except (ValueError, TypeError):
        pass  # Mantener como string si la conversión falla
    # Mapeo de operadores (usando el módulo operator)
    ops = {
        '==': operator.eq,
        '!=': operator.ne,
        '>': operator.gt,
        '<': operator.lt,
        '>=': operator.ge,
        '<=': operator.le,
        'contains': lambda col, val: col.astype(str).str.contains(str(val), case=case_sensitive),
        'not_contains': lambda col, val: ~col.astype(str).str.contains(str(val), case=case_sensitive)
    }
    # Verificar si el operador es válido
    if operator_str not in ops:
        return jsonify({"error": "Operador no válido"}), 400
    # Aplicar el filtro
    filtered_df = df[ops[operator_str](df[column], value)]
    filtered_df_json = convertir_a_json(filtered_df)
    return filtered_df_json

@app.route("/filtrar_filas", methods=['POST'])
def filtrar_filas():
    try:
        data = request.get_json()
        filtered_df_json = eliminar_filas( data )
        nueva_tabla = json.dumps( filtered_df_json )
        return Response(nueva_tabla, mimetype='application/json')
    except KeyError as e:
        return jsonify({"error": f"Columna no encontrada: {str(e)}"}), 400
    except Exception as e:
        return jsonify({"error": f"Error interno: {str(e)}"}), 500

@app.route("/plantillas_kpi", methods=['GET','POST'])
def plantillas_kpi():
    global conn
    gvar = get_global_variables()
    if gvar["username"] == "":
        return autenticar()
    # Obtener datos usuario de la sesion
    query = text('SELECT * FROM usuarios WHERE username="'+gvar['username']+'" AND password="'+gvar['password']+'"')
    try:
        user = conn.execute(query).fetchone()
    except SQLAlchemyError as e:
        conn = mydb.connect()
        user = conn.execute(query).fetchone()
    if request.method == 'POST':
        # Leemos la plantilla  KPI y lo salvamos en el servidor...
        plantillaKPI = request.files["plantilla_kpi"]
        # Leer denominacion y clase
        denominacion = request.form['denominacion']
        clase = request.form['clase']
        # Chequeamos si existe una plantilla en la base de datos con la misma denominacion...
        query = text('SELECT * FROM plantillas_kpi WHERE Denominacion="'+denominacion+'"')
        res = conn.execute(query).fetchall()
        if res is None or len(res):
            return jsonify({"Status":"ERROR", 
                            "Message":"Ya existe una plantilla con la misma denominación."})
        ext_doc = plantillaKPI.filename.rsplit('.', 1)[1].lower()
        # =========== Salvar a base de datos y adquirimos el id para nombrar la plantilla ======================
        # salvar a la tabla plantillas_kpi
        plantilla_kpi = metadata.tables['plantillas_kpi']
        ins = plantilla_kpi.insert().values(idUsuario=user.idUsuario,Denominacion=denominacion,Clase=clase)
        result = conn.execute(ins)
        conn.commit()
        id_plantilla = result.lastrowid
        if id_plantilla:
            url_kpi = 'templates/kpi/'+str(user.idUsuario)+'_'+str(id_plantilla)+'.' + ext_doc
            plantillaKPI.save(url_kpi)
            upd = plantilla_kpi.update().where(plantilla_kpi.c.idPlantillaKPI==id_plantilla).values(archivo_kpi=url_kpi)
            result = conn.execute(upd)
            conn.commit()
            # Enviamos la fila de tabla a la lista de las plantillas perteneciente al usuario actual,
            # cuantas plantillas hay?
            query = text("SELECT * FROM plantillas_kpi")
            res = conn.execute(query).fetchall()
            if len(res) % 2 == 0: # es par!
                color = '#FFFFFF'
            else:
                color = '#CCCCCC'
            fila_plantilla_nueva = '<tr bgcolor="'+color+'" style="cursor:pointer" onmousedown="Cargar_Plantilla(\''+url_kpi+'\')"><td>'+denominacion+'</td><td>'+clase+'</td></tr>'
            return jsonify({"Status":"OK","Fila_nueva": Markup(fila_plantilla_nueva)})
        return jsonify({"Status":"ERROR",
                        "Message":"No se pudo archivar la plantilla actual."})
    else: # GET
        # Aqui, antes de enviar la renderizacion de la plantilla, conformamos las dos listas de estas
        # que son las de otros y las propias del usuario actual
        # ==== Primero las del usuario actual...
        Lista_plantillas_propias = ''
        query = text('SELECT * FROM plantillas_kpi WHERE idUsuario='+str(user.idUsuario))
        res = conn.execute(query).fetchall()
        if len(res):
            alternar = True
            for row in res:
                alternar = not alternar
                if alternar:
                    color = '#CCCCCC'
                else:
                    color = '#FFFFFF'
                Lista_plantillas_propias += '<tr bgcolor="'+color+'" style="cursor:pointer" onmousedown="Cargar_Plantilla(\''+row.archivo_kpi+'\')"><td>'+row.Denominacion+'</td><td>'+row.Clase+'</td></tr>'
        Lista_plantillas_otros = ''
        query = text('SELECT * FROM plantillas_kpi WHERE idUsuario<>'+str(user.idUsuario))
        res = conn.execute(query).fetchall()
        if len(res):
            alternar = True
            for row in res:
                alternar = not alternar
                if alternar:
                    color = '#CCCCCC'
                else:
                    color = '#FFFFFF'
                Lista_plantillas_otros += '<tr bgcolor="'+color+'" style="cursor:pointer" onmousedown="Cargar_Plantilla(\''+row.archivo_kpi+'\')"><td>'+row.Denominacion+'</td><td>'+row.Clase+'</td></tr>'
        return render_template("plantillas_kpi.html", Lista_plantillas_otros=Markup(Lista_plantillas_otros), Start=False,
                               Lista_plantillas_propias=Markup(Lista_plantillas_propias), **gvar)
    
@app.route("/obtener_misplantillaskpi", methods=['GET'])
def obtener_misplantillaskpi():
    global conn
    gvar = get_global_variables()
    # Obtener datos usuario de la sesion
    query = text('SELECT * FROM usuarios WHERE username="'+gvar['username']+'" AND password="'+gvar['password']+'"')
    try:
        user = conn.execute(query).fetchone()
    except SQLAlchemyError as e:
        conn = mydb.connect()
        user = conn.execute(query).fetchone()
    # Obtener las plantillas KPI del usuario actual
    query = text('SELECT * FROM plantillas_kpi WHERE idUsuario=' + str(user.idUsuario))
    plantillas = conn.execute(query).fetchall()
    # Formatear la lista de plantillas
    lista_plantillas = [{"id": p.idPlantillaKPI, "denominacion": p.Denominacion, "archivo": p.archivo_kpi} for p in plantillas]
    return jsonify(lista_plantillas)

@app.route("/obtener_otrasplantillaskpi", methods=['GET'])
def obtener_otrasplantillaskpi():
    global conn
    gvar = get_global_variables()
    # Obtener datos usuario de la sesion
    query = text('SELECT * FROM usuarios WHERE username="'+gvar['username']+'" AND password="'+gvar['password']+'"')
    try:
        user = conn.execute(query).fetchone()
    except SQLAlchemyError as e:
        conn = mydb.connect()
        user = conn.execute(query).fetchone()
    # Obtener todas las plantillas KPI que no son del usuario actual...
    query = text('SELECT * FROM plantillas_kpi WHERE idUsuario<>' + str(user.idUsuario))
    plantillas = conn.execute(query).fetchall()
    # Formatear la lista de plantillas
    lista_plantillas = [{"id": p.idPlantillaKPI, "denominacion": p.Denominacion, "archivo": p.archivo_kpi} for p in plantillas]
    return jsonify(lista_plantillas)

import re
# Variable global que guarda el patron para hallar campos en las plantillas
patron_CAMPO = r'__([a-zA-Z0-9_-]+)__'
def extraer_campos_plantilla(path_plantilla):
    """
    Extrae campos de una plantilla de Word que siguen la sintaxis __CAMPO__,
    buscando en todos los elementos del documento (párrafos, tablas, celdas).
    
    Parámetros:
        path_plantilla (str): Ruta al archivo Word (.docx).
    Retorna:
        list: Lista de campos únicos encontrados en la plantilla.
    """
    global patron_CAMPO
    doc = Document(path_plantilla)
    campos = set()
    # Buscar en párrafos normales
    for paragraph in doc.paragraphs:
        if matches := re.findall(patron_CAMPO, paragraph.text):
            campos.update(matches)
    # Buscar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if matches := re.findall(patron_CAMPO, paragraph.text):
                        campos.update(matches)
    # Opcional: Buscar en encabezados y pies de página
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if matches := re.findall(patron_CAMPO, paragraph.text):
                campos.update(matches)
        for paragraph in section.footer.paragraphs:
            if matches := re.findall(patron_CAMPO, paragraph.text):
                campos.update(matches)
    # Iterar sobre todos los elementos de dibujo (shapes) en el documento
    for element in doc.element.body.iter():
        if etree.QName(element.tag).localname == 'drawing':
            # Buscar cuadros de texto dentro del elemento drawing
            for txbx in element.iter("{*}txbxContent"):
                for paragraph in txbx.iter("{*}p"):
                    text = ''.join(node.text for node in paragraph.iter("{*}t") if node.text)
                    if "__" in text:  # Buscar placeholders
                        if matches := re.findall(patron_CAMPO, text.strip()):
                            campos.update(matches)
    return jsonify(list(campos))

idx_tmp = 0 # indice de documentos temporales que se can generando
@app.route("/chequear_campos_plantilla", methods=['POST'])
def chequear_campos_plantilla():
    global idx_tmp
    # Obtener el archivo de la plantilla seleccionada
    documento = request.files["plantilla_kpi"]
    ext_doc = documento.filename.rsplit('.', 1)[1].lower()
    idx_tmp += 1
    tmp_doc = 'static/tmp/tmp'+str(idx_tmp)+'.' + ext_doc
    documento.save(tmp_doc)
    return extraer_campos_plantilla(tmp_doc)

@app.route("/leer_campos_plantilla", methods=['POST'])
def leer_campos_plantilla():
    # Obtener el archivo de la plantilla seleccionada
    archivo_plantilla = request.form["archivo_plantilla"]
    path_plantilla = os.path.join(app.root_path, archivo_plantilla)
    # Leer el archivo Word y extraer los campos con formato {{CAMPO}}
    return extraer_campos_plantilla(path_plantilla)

@app.route("/administrar")
def administrar():
    global conn
    gvar = get_global_variables()
    # Obtener listado de los usuarios que no están aprobados...
    query = text("SELECT * FROM usuarios WHERE aprobado = 0")
    try:
        result = conn.execute(query).fetchall()
    except SQLAlchemyError as e:
        conn = mydb.connect()
        result = conn.execute(query).fetchall()
    if len(result) > 0:
        listado_noaprobados = '<tr>'
        listado_noaprobados += '<th width="3%" bgcolor="#9999FF" scope="col"><div align="left">No.</div></th>'
        listado_noaprobados += '<th width="45%" bgcolor="#9999FF" scope="col"><div align="left">Nombres y Apellidos</div></th>'
        listado_noaprobados += '<th width="10%" bgcolor="#9999FF" scope="col"><div align="left">Tipo o Rol</div></th>'
        listado_noaprobados += '<th width="24%" bgcolor="#9999FF" scope="col"><div align="left">Email</div></th>'
        listado_noaprobados += '<th width="11%" bgcolor="#9999FF" scope="col"><div align="left">Foto</div></th>'
        listado_noaprobados += '<th width="7%" bgcolor="#9999FF" scope="col"><div align="left">Decisi&oacute;n</div></th>'
        listado_noaprobados += '</tr>'
        i = 1
        for row in result:
            fila = '<tr>'
            fila += '<td>' + str(i) + '</td>'
            fila += '<td>' + row.nom_ape + '</td>'
            fila += '<td>' + row.rol + '</td>'
            fila += '<td>' + row.email + '</td>'
            if row.url_foto != "" and row.url_foto is not None:
                url_foto = 'fotos_usuarios/'+row.url_foto
            else:
                url_foto = 'img/anonimo_black.png'
            fila += '<td><img height="40px" src="'+url_for('static', filename=url_foto)+'"/></td>'
            fila += '<td><a href="/aprobar/'+str(row.idUsuario)+'">Aprobar</a> o <a href="/desaprobar/'+str(row.idUsuario)+'">Borrar</a></td>'
            fila += '</tr>'
            listado_noaprobados += fila
            i += 1
    else:
        listado_noaprobados = '<div class="sin_pendientes"><h3><em>No hay usuarios pendientes de aprobación para ser registrados al sistema.</em></h3></div>'
    return render_template("administrar.html", Start=False, listado_noaprobados = Markup(listado_noaprobados), **gvar)

# Aprobar a usuario existente ante peticion de registrarse
@app.route("/aprobar/<int:usuario_id>")
def aprobar_usuario(usuario_id = 0):
    global conn
    if usuario_id:
        usuarios = metadata.tables['usuarios']
        oper = usuarios.update().where(usuarios.c.idUsuario == usuario_id).values(aprobado = True)
        try:
            conn.execute(oper)
        except SQLAlchemyError as e:
            conn = mydb.connect()
            conn.execute(oper)
        conn.commit()
    return administrar()

# Borrar a usuario existente ante peticion de registrarse
@app.route("/desaprobar/<int:usuario_id>")
def desaprobar_usuario(usuario_id = 0):
    global conn
    if usuario_id:
        query = text("SELECT * FROM usuarios WHERE idUsuario="+str(usuario_id))
        try:
            res = conn.execute(query).fetchone()
        except SQLAlchemyError as e:
            conn = mydb.connect()
            res = conn.execute(query).fetchone()
        if res is not None:
            # borramos primero la foto que está en el campo url_foto...
            path_photo = os.path.join(app.root_path, 'static', 'fotos_usuarios', res.url_foto)
            if os.path.exists(path_photo):
                os.remove(path_photo)
            # ... ahora borramos el registro de la tabla usuarios
            query = text("DELETE FROM usuarios WHERE idUsuario="+str(usuario_id))
            conn.execute(query)
            conn.commit()
    return administrar()

IndexUnicoInformeKPI = 0 # Variable que controla la unicidad de los informes KPI generados
#... para dos o mas usuarios conectados al servidor, haciendo el mismo objetivo.
# Tabla de saltos a las funciones de clasificacion
TablaLookUpClasif = {
    "Clasif_Arbol_decision" : Clasif_ArbolDecision,
    "Clasif_Random_forest" : Clasif_Randomforest,
    "Clasif_Regresion_logistica" : Clasif_Regresionlogistica,
    "Clasif_SVM" : Clasif_SVM
}
# onteo de tablas para su ID unico
count_t = 1000
@app.route("/ejecutar_kpi", methods=['POST', 'GET'])
def ejecutar_kpi():
    global conn, IndexUnicoInformeKPI
    gvar = get_global_variables()
    if gvar["username"] == "":
        return autenticar()
    if request.method == 'POST': # ======================================
        temporales_doc = []
        # Verificar si se envió un archivo o una URL
        if 'documento' in request.files:
            # Leer el documento cargado y salvarlo temporalmente en el servidor
            documento = request.files['documento']
            ext_doc = os.path.splitext(documento.filename)[1].lower()
            dest_folder = 'static/tmp'
            tmp_doc = os.path.join(dest_folder, 'tmp' + ext_doc)
            documento.save(tmp_doc)
            if not ext_doc in ['.zip', '.rar']:    
                temporales_doc.append(tmp_doc)
            else: # se envio un doc zip o rar
                if ext_doc == '.zip':
                    with zipfile.ZipFile(tmp_doc, 'r') as zip_ref:
                        # Listar archivos en el ZIP
                        archivos = zip_ref.namelist()
                        for idx, archivo in enumerate(archivos, start=1):
                            # Obtener la extensión del archivo original
                            nombre_original = os.path.basename(archivo)
                            extension = os.path.splitext(nombre_original)[1]  # Ej: ".docx"                            
                            # Nuevo nombre: tmp1<ext>, tmp2<ext>, ...
                            nuevo_nombre = f"tmp{idx}{extension}"
                            ruta_destino = os.path.join(dest_folder, nuevo_nombre)
                            # Extraer y renombrar el archivo
                            with zip_ref.open(archivo) as file_in_zip:
                                with open(ruta_destino, 'wb') as file_out:
                                    file_out.write(file_in_zip.read())
                            print(f"Extraído: {archivo} -> {nuevo_nombre}")
                            temporales_doc.append(ruta_destino)
                elif ext_doc == '.rar':
                    with rarfile.RarFile(tmp_doc, 'r') as rar_ref:
                        for idx, archivo in enumerate(rar_ref.infolist(), start=1):
                            # Obtener nombre base y extensión del archivo
                            nombre_base = os.path.basename(archivo.filename)
                            nombre, extension = os.path.splitext(nombre_base)
                            # Saltar directorios (si los hay)
                            if not nombre_base:
                                continue
                            # Nuevo nombre: tmp1<ext>, tmp2<ext>, etc.
                            nuevo_nombre = f"tmp{idx}{extension}"
                            ruta_completa = os.path.join(dest_folder, nuevo_nombre)
                            # Extraer y renombrar el archivo
                            with rar_ref.open(archivo) as archivo_rar:
                                with open(ruta_completa, 'wb') as archivo_destino:
                                    archivo_destino.write(archivo_rar.read())
                            print(f"Extraído: {archivo.filename} -> {nuevo_nombre}")
                            temporales_doc.append(ruta_completa)
        elif 'url' in request.form:
            # Leer el documento desde la URL
            url = request.form['url']
            try:
                respuesta = requests.get(url)
                if respuesta.status_code == 200:
                    ext_doc = os.path.splitext(nombre_base)[1].lower()
                    tmp_doc = dest_folder + ext_doc
                    with open(tmp_doc, 'wb') as archivo:
                        archivo.write(respuesta.content)        
                    temporales_doc.append( tmp_doc )            
            except requests.exceptions.RequestException as e:
                return jsonify({"Status": "Error", "Message": f"Error en la solicitud: {str(e)}"})
        else:
            return jsonify({"Status": "Error", "Message": "No se proporcionó un archivo ni una URL."})
        
        # Obtener el TimeZone del cliente  
        client_timezone = request.form["timezone"]
        # Tenemos el o los documento(s) y su extension: leemos ahora sus tablas...
        lista_tablas_doc = []
        for tmp_doc in temporales_doc:
            ext_doc = os.path.splitext(tmp_doc)[1].lower()
            lista_tablas_doc.append(json.loads(dar_tablas(tmp_doc, ext_doc)))
        # en session esta guardado workflow_json...
        path_json = session.get('path_json', None)
        if path_json is None:
            return jsonify({"Status": "Error", "Message": "No se halla archivo JSON!"})
        workflow_json, conecciones = leer_json_workflow(path_json)
        # COMIENZA LA EJECUCION DEL MODELO...
        # 1.- Chequear el Matcheo entre las tablas halladas en el documento de entrada 
        # y las que tiene el objeto Documento en workflow_json...
        # Diccionario de acciones
        def acc_documento(elem):
            nonlocal lista_tablas_doc, workflow_json
            # Buscar las tablas por documento y Verificar 
            # si se tienen el mismo numero de tablas ...
            for tablas_doc in lista_tablas_doc:
                if len(elem["tablas"]) == len(tablas_doc):
                    # Se asume que no habra errores de matcheo entre las tablas a comparar
                    error_matcheo = 0 
                    i = 0
                    for tablad,tablam in zip(tablas_doc, elem["tablas"]):
                        columns = list(tablad.keys())
                        if columns != tablam["columnas"]:
                            error_matcheo += 1 # Se van sumando el numero de matcheos erroneos
                        else:
                            # Actualizar la tabla en al workflow_json...
                            elem["tablas"][i]["table"] = tablad
                            elem["tablas"][i]["rows"] = len(tablad[columns[0]])
                            elem["tablas"][i]["cols"] = len(columns)
                        i += 1
                    if not error_matcheo:
                        return (True, "")
                    else:
                        continue
            return (False,"No se haya un documento de entrada con el que necesita el modelo KPI.")
        
        def acc_limpieza(elem):
            nonlocal workflow_json, conecciones
            # Hallar el los elementos conectados a este elemento Limpieza
            idFroms = [d["from"] for d in conecciones if d["to"] == elem["id"]]
            if not idFroms:
                return (False, "El obj Limpieza no esta conectado desde algun obj.")
            objFrom = None
            for obj in workflow_json:
                if obj["id"] in idFroms and obj["tipo"] in ["Documento", "Limpieza", "Conversion"]:
                    if obj["tipo"] == "Documento":
                        lista_tablas = obj["tablas"]
                    else:
                        lista_tablas = obj["resultado"]
                    tablaStruc = next((tab for tab in lista_tablas if tab["idTable"] == elem["idtable"]), None)
                    if not tablaStruc is None:
                        objFrom = obj
                        break
            if objFrom is None:
                return (False, "La tabla referenciada en obj Limpieza no aparece en los objetos conectados.")
            match(elem["proceso"]):
                case "Quitar_duplicados":
                    columnas_ref = elem["cols"]
                    tabla = tablaStruc["table"]
                    # Aplicar Quitar duplicados
                    nueva_tabla = convertir_a_json(Quitar_duplicados(tabla, columnas_ref))
                    elem["resultado"][0]["table"] = nueva_tabla
                    return (True, "OK")
                case "Recuperar_num_corruptos":
                    columnas_ref = elem["cols"]
                    tabla = tablaStruc["table"]
                    # Aplicar Quitar duplicados
                    nueva_tabla = convertir_a_json(RecuperarNumCorruptos(tabla, columnas_ref))
                    elem["resultado"][0]["table"] = nueva_tabla
                    return (True, "")
                case "Eliminar_filasfcond":
                    data = {
                        'table' : tablaStruc["table"],
                        'column' : elem["cols"],
                        'operator' : elem["operator"],
                        'value' : elem["value"]
                    }
                    nueva_tabla = eliminar_filas( data )
                    elem["resultado"][0]["table"] = nueva_tabla
                    return (True,"")
            return (False,"No se halló proceso para objeto Limpieza.")
        
        def acc_conversion(elem):
            nonlocal workflow_json, conecciones
            # Hallar el los elementos conectados a este elemento Limpieza
            idFroms = [d["from"] for d in conecciones if d["to"] == elem["id"]]
            if not idFroms:
                return (False, "El obj Conversión no está conectado desde algun obj.")
            objFrom = None
            for obj in workflow_json:
                if obj["id"] in idFroms and obj["tipo"] in ["Limpieza", "Conversion"]:
                    if obj["tipo"] == "Documento":
                        lista_tablas = obj["tablas"]
                    else:
                        lista_tablas = obj["resultado"]
                    tablaStruc = next((tab for tab in lista_tablas if tab["idTable"] == elem["idtable"]), None)
                    if not tablaStruc is None:
                        objFrom = obj
                        break
            if objFrom is None:
                return (False, "La tabla referenciada en obj Conversión no aparece en los objetos conectados.")
            proceso = elem["proceso"]
            match(proceso):
                case "Agrupar_filas":
                    data = {
                        'table' : tablaStruc["table"],
                        'groupBy' : elem["cols"],
                        'operations' : elem["operations"]
                    }
                    nueva_tabla = agrupar_filas( data )
                    elem["resultado"][0]["table"] = nueva_tabla
                    return (True,"")
                case "Ordenar_filas":
                    data = {
                        'table' : tablaStruc["table"],
                        'sort_by' : elem["cols"],
                        'ascending' : elem["ascending"]
                    }
                    nueva_tabla = ordenar_filas( data )
                    elem["resultado"][0]["table"] = nueva_tabla
                    return (True,"")
            return (False,"No se halló proceso para objeto Conversión.")
        
        def acc_procesamiento(elem):
            nonlocal workflow_json, conecciones
            # Hallar el elemento tipo Limpieza/Conversion al cual se conecta este elemento Procesamiento
            idLimCon = next((dicc["from"] for dicc in conecciones if dicc["to"] == elem["id"]), None)
            if idLimCon is None:
                return (False, "El obj Procesamiento no esta conectado desde algun obj Limpieza/Conversion.")
            objLimCon = next((obj for obj in workflow_json if obj["id"]==idLimCon), None)
            if objLimCon is None:
                return (False, "El obj Procesamiento no esta conectado desde algun obj Limpieza/Conversion.")
            idtable = elem["idtable"] # id tabla a aplicarle el Procedimiento
            columnas = elem["cols"] # columnas requeridas 
            proceso = elem["proceso"]
            match(proceso):
                case "Prediccion_lineal":
                    for tablaid in objLimCon["resultado"]:
                        if tablaid["idTable"]==idtable:
                            table = tablaid["table"]
                            xCol = columnas[0]
                            yCol = columnas[1]
                            x = table[xCol]
                            y = table[yCol]
                            epocas = elem["epocas"]
                            df = Prediccion_lineal(x, y, xCol, yCol, epocas)
                            nueva_tabla = convertir_a_json( df )
                            elem["resultado"][0]["table"] = nueva_tabla
                            elem["resultado"][0]["rows"] = len(nueva_tabla[xCol])
                            return (True,"")
                    return (False,"No se hallo la tabla requerida para el Procesamiento.")    
                case "Prediccion_nolineal":
                    for tablaid in objLimCon["resultado"]:
                        if tablaid["idTable"]==idtable:
                            table = tablaid["table"]
                            xCol = columnas[0]
                            yCol = columnas[1]
                            x = table[xCol]
                            y = table[yCol]
                            epocas = elem["epocas"]
                            df = Prediccion_nolineal(x, y, xCol, yCol, epocas)
                            nueva_tabla = convertir_a_json( df )
                            elem["resultado"][0]["table"] = nueva_tabla
                            elem["resultado"][0]["rows"] = len(nueva_tabla[xCol])
                            return (True,"")
                    return (False,"No se hallo la tabla requerida para el Procesamiento.")
                case "Prediccion_SeriesTemporales":
                    for tablaid in objLimCon["resultado"]:
                        if tablaid["idTable"]==idtable:
                            table = tablaid["table"]
                            xCol = columnas[0]
                            yCol = columnas[1]
                            epocas = elem["epocas"]
                            df = pd.DataFrame(table)
                            df = Prediccion_SeriesTemporales(df, xCol, yCol, epocas)
                            nueva_tabla = convertir_a_json( df )
                            elem["resultado"][0]["table"] = nueva_tabla
                            elem["resultado"][0]["rows"] = len(nueva_tabla[xCol])
                            return (True,"")
                    return (False,"No se hallo la tabla requerida para el Procesamiento.")
                case "Estadistica_descriptiva":
                    for tablaid in objLimCon["resultado"]:
                        if tablaid["idTable"]==idtable:
                            table = tablaid['table']
                            df = EstadisticaDescriptiva(table, columnas)
                            # Convertir el DataFrame de nuevo a un diccionario
                            nueva_tabla = convertir_a_json(df)
                            nuevas_cols = list(nueva_tabla.keys())
                            elem["resultado"][0]["table"] = nueva_tabla
                            elem["resultado"][0]["cols"] = len(nuevas_cols)
                            elem["resultado"][0]["rows"] = len(nueva_tabla[nuevas_cols[0]])
                            elem["resultado"][0]["columnas"] = nuevas_cols
                            return (True,"")
                case "Clasif_Arbol_decision" | "Clasif_Random_forest" | \
                     "Clasif_Regresion_logistica" | "Clasif_SVM":
                    for tablaid in objLimCon["resultado"]:
                        if tablaid["idTable"]==idtable:
                            table = tablaid["table"]
                            colClasif = elem["colClasif"]
                            fun_clasif = TablaLookUpClasif.get(proceso, Clasif_default)
                            resultado = fun_clasif(table, columnas, colClasif)
                            for clave,table in resultado.items():
                                tabla = next((obj for obj in elem["resultado"] if obj["descrip"]==clave), None)
                                tabla["table"] = table
                            return (True,"")
                    return (False,"No se hallo la tabla requerida para el Procesamiento.")
            return (False,"El proceso dado para Procesamiento no existe.")

        def acc_grafico(elem):
            nonlocal workflow_json, conecciones
            type = elem["type"] # Tipo de grafico
            # Hallar el elemento tipo Limpieza/Procesamiento al cual se conecta este elemento Grafico
            idLimProc = next((dicc["from"] for dicc in conecciones if dicc["to"] == elem["id"]), None)
            if idLimProc is None:
                return (False, "El obj Grafico no esta conectado desde algun obj de Limpieza p Procesamiento.")
            objLimProc = next((obj for obj in workflow_json if obj["id"] == idLimProc), None)
            if objLimProc is None:
                return (False, "El obj Grafico no esta conectado desde algun obj de Limpieza o Procesamiento.")
            idtable = elem["idtable"] # id tabla a aplicarle el Grafico
            tabla = objLimProc["resultado"][0]
            if tabla["idTable"] != idtable:
                return (False,"No se ha hallado la tabla para hacer graficado.")
            columnas = elem["cols"]
            colX = columnas[0]
            colY = columnas[1]
            table = tabla["table"]
            # Extraer los datos de las columnas X e Y
            datosX = table[colX]
            datosY = table[colY]
            match(type):
                case "Generar_grafico_lineas" | "Generar_grafico_barras" | \
                    "Generar_grafico_dispersion" | "Generar_grafico_pastel" | \
                    "Generar_grafico_histograma":
                    # Verificar si en el obj grafico se an~adieron otros datos a plotear
                    if "type_prev" in elem.keys():
                        type_prev = elem["type_prev"]
                        fun_graf = accTipoGraf.get(type_prev, accion_grafdefault)
                        idtable_prev = elem["idtable_prev"]
                        # Quiere decir que objLimProc es un obj Procesamiento, buscar entonces el obj anterior 
                        # de donde procede la tabla con id idtable2...
                        idPrev = next((dicc["from"] for dicc in conecciones if dicc["to"] == objLimProc["id"]), None)
                        objPrev = next((obj for obj in workflow_json if obj["id"] == idPrev), None)
                        if objPrev is not None:
                            tablaPrev = objPrev["resultado"][0]
                            if tablaPrev["idTable"] == idtable_prev:
                                datosX2 = tablaPrev["table"][colX]
                                datosY2 = tablaPrev["table"][colY]
                                # Graficado doble
                                (buf, msg) = fun_graf(datosX2, datosY2, colX, colY, 
                                                    datosX, datosY, type)
                            else:
                                return (False, "Tabla id no encontrada!")
                    else:
                        # Graficado sencillo
                        fun_graf = accTipoGraf.get(type, accion_grafdefault)
                        (buf, msg) = fun_graf(datosX, datosY, colX, colY)
                    if buf is None:
                        return (False, msg)
                    else:
                        elem["resultado"] = buf
                        return (True,"OK")
                case _:
                    return (False,"Tipo de grafico no identificado.")
        
        def acc_informekpi(elem):
            nonlocal workflow_json, conecciones
            global patron_CAMPO
            # Insertar una tabla en el parrafo
            def insert_table(location, table_data, is_paragraph=True):
                """Inserta una tabla en el documento sin borrar otros elementos.
                Argumentos:
                    location: Párrafo (si is_paragraph=True) o celda (si is_paragraph=False).
                    table_data: Diccionario {nombre_columna: [valores]}.
                    is_paragraph: Booleano que indica si location es un párrafo.
                Retorna:
                    La tabla insertada o None si hay error.
                """
                try:
                    # Obtener el documento
                    doc = location.part.document
                    # Crear tabla con bordes (similar al original)
                    table = doc.add_table(rows=1, cols=len(table_data))
                    tbl = table._tbl
                    # Configurar bordes (opcional, igual al original)
                    borders = OxmlElement('w:tblBorders')
                    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        element = OxmlElement(f'w:{border}')
                        element.set(qn('w:val'), 'single')
                        element.set(qn('w:sz'), '4')
                        element.set(qn('w:color'), '000000')
                        borders.append(element)
                    tbl.tblPr.append(borders)
                    # Llenar encabezados (en negrita)
                    headers = list(table_data.keys())
                    for i, header in enumerate(headers):
                        table.cell(0, i).text = header
                        for paragraph in table.cell(0, i).paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    # Llenar filas
                    for row_idx in range(len(table_data[headers[0]])):
                        row_cells = table.add_row().cells
                        for col_idx, header in enumerate(headers):
                            row_cells[col_idx].text = str(table_data[header][row_idx])
                    # Insertar en la ubicación correcta
                    if is_paragraph:
                        # Insertar después del párrafo (sin borrar nada)
                        paragraph = location
                        paragraph.text = ""
                        paragraph._p.addnext(tbl)
                    else:
                        # Insertar dentro de la celda (reemplazando contenido)
                        cell = location
                        cell.text = ""
                        cell._tc.append(tbl)
                    return table
                except Exception as e:
                    logging.error(f"Error insertando tabla: {str(e)}")
                    if is_paragraph:
                        location.text = f"[ERROR TABLA: {str(e)}]"
                    return None

            # Añade esto ANTES de la función acc_informekpi
            def _get_cell_width_twips(paragraph):
                """Obtiene el ancho REAL de celda en twips (1/1440 de pulgada)"""
                if hasattr(paragraph, '_parent') and isinstance(paragraph._parent, _Cell):
                    return paragraph._parent.width - 360  # Restar márgenes (0.25")
                return 6 * 1440  # Default 6 pulgadas si no es tabla

            def insertar_imagen_corregida(paragraph, buf, is_paragraph):
                try:
                    # Validar buffer
                    if not isinstance(buf, io.BytesIO):
                        buf = io.BytesIO(buf) if isinstance(buf, bytes) else None
                        if not buf:
                            raise ValueError("Buffer de imagen no válido")
                    buf.seek(0)
                    # Calcular dimensiones
                    with Image.open(buf) as img:
                        img_width_px, img_height_px = img.size
                        dpi = img.info.get('dpi', (96, 96))[0]
                        width_in = img_width_px / dpi
                        height_in = img_height_px / dpi
                        aspect_ratio = height_in / width_in
                        # Ajustar al ancho disponible
                        available_twips = _get_cell_width_twips(paragraph)
                        available_in = available_twips / 1440
                        max_width_in = 6.5 # Ancho máximo permitido
                        new_width = min(available_in * 0.9, width_in, max_width_in)
                        new_height = new_width * aspect_ratio
                        # Insertar imagen
                        paragraph.text = "" # Limpiar texto existente (placeholder)
                        if not is_paragraph: # Es una celda de tabla
                            para = paragraph.paragraphs[0] if paragraph.paragraphs else paragraph.add_paragraph()
                            run = para.add_run()
                            run.add_picture(
                                buf,
                                width=Inches(new_width),
                                height=Inches(new_height)
                            )
                        else:
                            run = paragraph.add_run()
                            run.add_picture(
                                buf,
                                width=Inches(new_width),
                                height=Inches(new_height)
                            )

                except Exception as e:
                    paragraph.text = f"[ERROR IMAGEN: {str(e)}]"

            def insertar_valor_campo(campo, paragraph, is_paragraph=True):
                nonlocal asociaciones
                global conn

                # Método universal para reemplazar
                def replace_in_paragraph(paragraph, old, new):
                    # Es un párrafo normal (tiene runs)
                    if hasattr(paragraph, 'runs'):
                        paragraph.text = paragraph.text.replace(old, new)
                    # Es XML crudo (cuadros de texto, tablas profundas)
                    else:
                        # Método 100% seguro para encontrar nodos de texto
                        for elem in paragraph.iter():
                            if elem.tag.endswith('}t'):  # Busca <w:t>
                                if elem.text and old in elem.text:
                                    elem.text = elem.text.replace(old, new)

                # continuación función insertar_valor_campo
                if campo not in asociaciones.keys():
                    return
                campo_formateado = f"__{campo}__"
                valor = asociaciones[campo]        
                if isinstance(valor, str):
                    # Manejo de strings (existente)
                    if valor == "NOM_APE_EJECUTOR":
                        query = text("SELECT * FROM usuarios WHERE username='"+gvar["username"]+"' AND password='"+gvar["password"]+"'")
                        try:
                            res = conn.execute(query).fetchone()
                        except SQLAlchemyError as e:
                            conn = mydb.connect()
                            res = conn.execute(query).fetchone()
                        Nom_Ape = res.nom_ape
                        replace_in_paragraph(paragraph, campo_formateado, Nom_Ape)
                    elif valor == "FECHA_HORA":
                        # Obtener hora actual en la zona horaria del cliente
                        client_time = datetime.now(pytz.timezone(client_timezone))
                        # Formatear (ej: "25/06/2024, 10:30 AM")
                        fecha_hora_formateada = client_time.strftime("%d/%m/%Y, %I:%M:%S %p") + " ("+client_timezone+")"
                        replace_in_paragraph(paragraph, campo_formateado, fecha_hora_formateada)
                elif isinstance(valor, dict):
                    tipo = valor["tipo"]
                    id = valor["id"]
                    obj = next((dicc for dicc in workflow_json if dicc["id"] == id), None)
                    if obj is None:
                        return (False, "No se halló un identificador de objeto para mostrar sus resultados en la plantilla KPI.")
                    elif tipo == "IMAGEN":
                        buf = obj["resultado"]
                        insertar_imagen_corregida(paragraph, buf, is_paragraph)
                    elif tipo == "TABLA":
                        try:
                            if "idx" in valor.keys():
                                idx = valor["idx"]
                            else:
                                idx = 0
                            tabla = obj["resultado"][idx]
                            table_data = tabla["table"]
                            insert_table(paragraph, table_data, is_paragraph)
                        except Exception as e:
                            logging.error(f"Error procesando tabla: {str(e)}")
                            paragraph.text = paragraph.text.replace(campo_formateado, f"[ERROR TABLA: {str(e)}]")
                return paragraph
        
            # --- CUERPO PRINCIPAL ---
            try:
                asociaciones = elem["asociaciones"]
                plantilla_kpi = elem["archivo"]
                doc = Document(plantilla_kpi)
                # Buscar en cada párrafo normal en el documento
                for paragraph in doc.paragraphs:
                    if matches := re.findall(patron_CAMPO, paragraph.text):
                        insertar_valor_campo(matches[0], paragraph)
                # Buscar en cada tabla
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            #for paragraph in cell.paragraphs:
                            if matches := re.findall(patron_CAMPO, cell.text):
                                insertar_valor_campo(matches[0], cell, is_paragraph=False)
                # Buscar en secciones de encabezados y pies de pagina
                for section in doc.sections:
                    # Buscar en encabezados
                    for paragraph in section.header.paragraphs:
                        if matches := re.findall(patron_CAMPO, paragraph.text):
                            insertar_valor_campo(matches[0], paragraph)
                    # Buscar en pies de pagina
                    for paragraph in section.footer.paragraphs:
                        if matches := re.findall(patron_CAMPO, paragraph.text):
                            insertar_valor_campo(matches[0], paragraph)
                # Iterar sobre todos los elementos de dibujo (shapes) en el documento
                for element in doc.element.body.iter():
                    if etree.QName(element.tag).localname == 'drawing':
                        # Buscar cuadros de texto dentro del elemento drawing
                        for txbx in element.iter("{*}txbxContent"):
                            for paragraph in txbx.iter("{*}p"):
                                texto = ''.join(node.text for node in paragraph.iter("{*}t") if node.text)
                                if "__" in texto:  # Buscar placeholders
                                    if matches := re.findall(patron_CAMPO, texto.strip()):
                                        insertar_valor_campo(matches[0], paragraph)
                return (True, doc)
            except Exception as e:
                logging.error(f"Error en acc_informekpi: {str(e)}")
                return (False, f"Error al procesar documento: {str(e)}")
        
        def accion_default(elem):
            return (False,"Hay un Objeto en el modelo que no está identificado.")

        acciones = {
            "Documento": acc_documento,
            "Limpieza": acc_limpieza,
            "Conversion" : acc_conversion,
            "Procesamiento": acc_procesamiento,
            "Grafico": acc_grafico,
            "InformeKPI": acc_informekpi
        }
        # Aquí se encuentra el corazon de la ejecucion del modelo para generar el informe KPI 
        # La estrategia es ir a cada obj Documento y siguiendo las conexiones, ejecutar cada objeto, 
        # hasta que nos encontremos, en cada hilo de conexión, al objeto InformeKPI, que debe ser único.
        # Luego se ejecuta dicho obj InformeKPI...
        def execute_next_conections( objPrime ):
            nonlocal workflow_json, conecciones
            accion = acciones.get(objPrime["tipo"], accion_default)
            Ok, resp = accion(objPrime)
            if not Ok:
                return jsonify({"Status": "Error", "Message": resp})
            # Hallar los elementos conectados al obj actual
            idTos = [d["to"] for d in conecciones if d["from"] == objPrime["id"]]
            for idTo in idTos:
                objTo = next((obj for obj in workflow_json if obj["id"] == idTo), None)
                if objTo["tipo"] != "InformeKPI":
                    execute_next_conections( objTo )
        
        # Ejecutando todos los obj desde los documentos...
        objDocs = [obj for obj in workflow_json if obj["tipo"] == "Documento"]
        for objDoc in objDocs:
            execute_next_conections( objDoc )
        # Despues, solo nos queda ejecutar al obj InformeKPI...
        objInfKPI = next((obj for obj in workflow_json if obj["tipo"] == "InformeKPI"), None)
        if objInfKPI is None:
            return jsonify({"Status": "Error", "Message": "Este modelo no tiene objeto Informe KPI !!! debe terminar su edición."})
        accion = acciones.get(objInfKPI["tipo"], accion_default)
        Ok, resp = accion(objInfKPI)
        if not Ok:
            return jsonify({"Status": "Error", "Message": resp})

        # Si llegamos aqui, en "resp" tenemos el documento actualizado
        # Vamos a generar un nombre generico usando la fecha y hora
        fecha_hora_actual = datetime.now()
        IndexUnicoInformeKPI += 1
        sufijo_unico_str = fecha_hora_actual.strftime("%d-%m-%Y_%H%M%S")+'_num'+str(IndexUnicoInformeKPI)
        # Generar url para la carpeta de temporales
        # En la parte final donde guardas los archivos:
        url_informe_kpi = 'static/tmp/informe_kpi_'+sufijo_unico_str
        url_informe_kpi_doc = url_informe_kpi + '.docx'
        url_informe_kpi_pdf = url_informe_kpi + '.pdf'
        try:
            # Guardar con manejo de errores
            resp.save(url_informe_kpi_doc)
            # Convertir a PDF
            pythoncom.CoInitialize()
            try:
                convert(url_informe_kpi_doc, url_informe_kpi_pdf)
            except Exception as e:
                logging.error(f"Error en conversión PDF: {str(e)}")
                return jsonify({
                    "Status": "Error", 
                    "Message": f"Error al generar PDF: {str(e)}",
                    "url_documento": url_informe_kpi_doc  # Enviar el DOCX como fallback
                })
            finally:
                pythoncom.CoUninitialize()
            return jsonify({
                "Status": "OK", 
                "url_documento": url_informe_kpi_pdf,
                "url_docx": url_informe_kpi_doc  # Opcional: para descarga del original
            })
        except Exception as e:
            logging.error(f"Error al guardar documento: {str(e)}")
            return jsonify({
                "Status": "Error",
                "Message": f"No se pudo guardar el documento: {str(e)}"
            })
    
    else:  # GET =========================================================
        # Listar todos los modelos definidos, que ya tengan un objeto InformeKPI
        query = text("SELECT * FROM modelos")
        try:
            res = conn.execute(query).fetchall()
        except SQLAlchemyError as e:
            conn = mydb.connect()
            res = conn.execute(query).fetchall()
        lista_documentos = ""
        primera_vez = True
        for row in res:
            # Leer el archivo modelo json...
            workflow_json, conections = leer_json_workflow(row.archivo_json)
            # Chequear al modelo, si tiene al menos una rama de ejecucion, 
            # desde un tipo Documento hasta un tipo InformeKPI...
            found_doc = False
            IniCon = {}
            for con in conections:
                objFrom = next((obj for obj in workflow_json if obj["id"] == con["from"]), None)
                if objFrom is not None and objFrom["tipo"] == "Documento":
                    found_doc = True
                    IniCon = con
                    break
            if not found_doc:
                continue
            # Hemos hallado un Documento, ahora vamos de conexion a conexion 
            # hasta hallar a un obj de tipo InformeKPI...
            found_kpi = False
            while not found_kpi and IniCon is not None:
                NextObj = next((obj for obj in workflow_json if obj["id"] == IniCon["to"]), None)
                if NextObj is None:
                    break
                if NextObj["tipo"] == "InformeKPI":
                    found_kpi = True
                else:
                    # pasar a la otra conexion...
                    IniCon = next((con for con in conections if con["from"] == IniCon["to"]), None)
            if not found_kpi:
                continue
            if primera_vez:
                primera_vez = False
                sel = 'selected="selected"'
            else:
                sel = ''
            lista_documentos += '<option value="' + row.archivo_json + '" '+sel+'>' + row.nombre + '</option>'
        return render_template("ejecutar_kpi.html", lista_documentos=Markup(lista_documentos), Start=False, **gvar)

def leer_json_workflow(path_json):
    # Leemos el archivo modelo json...
    with open(path_json, 'r', encoding='utf-8') as file:
        data_json = json.load(file)
    # Extraemos el archivo Plantilla KPI desde el primer elemento de data, atributo 
    # tipo == InformeKPI...
    listado = json.loads(data_json)
    return listado[0], listado[1]

# Entrada para obtener los tooltips de cada documento dado un modelo json...
@app.route('/get_tooltips', methods=['GET'])
def get_tooltips():
    # Obtener el nombre del archivo JSON desde la solicitud
    path_json = request.args.get('archivo_json')
    if not path_json:
        return jsonify({"Status": "Error", "Message": "No se proporcionó el nombre del archivo."})
    # Leer el archivo modelo json...
    workflow_json, conecc = leer_json_workflow(path_json)
    # Crear un diccionario para almacenar los tooltips
    tooltips = []
    # Recorrer los elementos del workflow_json
    for elem in workflow_json:
        if elem["tipo"] == "Documento":
            tooltips.append( elem["tooltip"] )
    return jsonify({"Status":"OK", "tooltips":tooltips})

@app.route('/convertir_word_pdf', methods=['POST'])
def convertir_word_a_pdf():
    global idx_tmp
    # Ruta del Word (enviada por el cliente)
    word_path = request.json.get('ruta_word')
    # Ruta de salida del PDF (en el servidor)
    idx_tmp += 1
    nombre_gen = 'doc_pdf_' + str(idx_tmp) + '.pdf' 
    pdf_path = os.path.join('static', 'tmp', nombre_gen)
    # Convertir Word a PDF (usando pythoncom en Windows)
    try:
        pythoncom.CoInitialize()  # Inicializar COM
        convert(word_path, pdf_path)  # Convertir
        pythoncom.CoUninitialize()  # Liberar recursos
    except Exception as e:
        return {"error": str(e)}, 500
    # Enviar el PDF al cliente (para abrir en nueva pestaña)
    return {
        "pdf_url": "/static/tmp/" + nombre_gen
    }

@app.route('/static/tmp/<filename>')
def servir_pdf(filename):
    return send_from_directory('static/tmp', filename)

# Ruta para servir archivos estáticos json (y extraer la plantilla Word asociada)
@app.route('/static/json/<filename>', methods=['GET'])
def servir_documento(filename):
    global idx_tmp
    # Leemos el archivo modelo json...
    path_json = 'static/json/' + filename
    workflow_json, conecc = leer_json_workflow(path_json)
    # Se guarda en session a workflow_json...
    session["path_json"] = path_json
    plantilla_kpi = ""
    # Buscar en workflow_json el objeto tipo InformeKPI...
    elem = next((elem for elem in workflow_json if elem["tipo"]=="InformeKPI"), None)
    if elem is None:
        return "El modelo seleccionado aun no tiene un objeto de tipo Informe KPI."        
    plantilla_kpi = elem["archivo"]
    if plantilla_kpi:
        # leemos el documento Word y lo convertimos a PDF 
        # luego lo enviamos al cliente...
        # Inicializar el sistema COM
        pythoncom.CoInitialize()
        try:
            idx_tmp += 1
            ruta = 'static/tmp'
            nombre_gen = 'tmp' + str(idx_tmp) + '.pdf'
            ruta_completa = ruta + '/' + nombre_gen
            convert(plantilla_kpi, ruta_completa)
        except Exception as e:
            print(f"Error durante la conversión: {e}")
        finally:
            # Liberar el sistema COM al finalizar
            pythoncom.CoUninitialize()
        return send_from_directory(ruta, nombre_gen)
    else:
        return Markup("Este modelo está incompleto, no tiene el <strong>Objeto InformeKPI</strong>. Por favor terminar de editarlo!.")

@app.route("/about")
def acerca_de():
    gvar = get_global_variables()
    return render_template("acerca_de.html", Start=False, **gvar)

# ------------- E N T R A D A   P R I N C I P A L   D E L   S E R V I D O R ----------------
if __name__=='__main__':
    app.run(debug = True)
